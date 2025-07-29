import streamlit as st
from docx import Document
import pandas as pd
import re
import yaml
import io
import docxedit
import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from streamlit_gsheets import GSheetsConnection
from docxtpl import DocxTemplate
from docx.shared import Inches, Pt
from openai import OpenAI
from st_aggrid import AgGrid, GridOptionsBuilder, JsCode
from modules.recommendations import *
# from modules.add_scores import *

##########################################################
st.set_page_config(
    page_title="Module 4",
    page_icon="📝",
    layout="centered",
    initial_sidebar_state="expanded",
)

##########################################################
# Set up OpenAI 
if 'behavior_observation_mod3' not in st.session_state:
    st.session_state.behavior_observation_mod3 = ""
if 'development_history_mod3' not in st.session_state:
    st.session_state.development_history_mod3 = ""

# Load OpenAI client 
client = OpenAI(api_key=st.secrets["openai_key"])

##################################################################
def transcribe_audio(audio_file, name='temp'):
    if audio_file:
        # Transcribe
        with st.spinner("Transcribing...", show_time=True):
            # result = whisper_model.transcribe(f"{name}.wav")
            result = client.audio.transcriptions.create(
                model="whisper-1", 
                file=audio_file, 
                response_format="text"
            )
        return result 

##########################################################
# Access Google Sheets

def get_abbreviation(test_name):
    # Split on en dash (\u2013)
    main_title = test_name.split('\u2013')[0].strip()
    
    # Split into words and get uppercase initials
    abbreviation = ''.join(word[0] for word in main_title.split() if word[0].isupper())
    
    return abbreviation

dropdowns = {}
connections = {}

# Create a connection object.
connections['All'] = st.connection(f"mod3_all", type=GSheetsConnection)
# Read object
df = connections['All'].read(
    ttl="30m",
    usecols=list(range(6)),
    nrows=30,
) 
for col_name in df.columns:
    dropdowns[col_name] = df[col_name].tolist()
    dropdowns[col_name] = [x for x in dropdowns[col_name] if str(x) != 'nan']

# DSM dropdowns
connections['DSM'] = st.connection(f"dsm", type=GSheetsConnection)
# Read object
df = connections['DSM'].read(
    ttl="30m",
    usecols=list(range(7)),
    nrows=15,
) 
for col_name in df.columns:
    dropdowns[col_name] = df[col_name].tolist()
    dropdowns[col_name] = [x for x in dropdowns[col_name] if str(x) != 'nan']
    dropdowns[col_name].append("None")


# Scores for sidebar
connections['Scores'] = st.connection(f"mod3_scores", type=GSheetsConnection)
# Read object
df = connections['Scores'].read(
    ttl="30m",
    usecols=list(range(6)),
    nrows=30,
) 
score_list = df.to_dict('records')

# Process data
scores = {}
check_scores = {}

for test in score_list:
    test_name = test["Test name"]
    abbr = get_abbreviation(test_name)
    scores[abbr] = {}

    scores[abbr]["Test name"] = test_name
    all_lines = []
    all_items = {}
    print(f"\nTest: {test_name}")
    
    for i in range(5):
        line_key = f"Line {i}"
        line_value = test.get(line_key)
        if line_value and str(line_value) != "nan":
            all_lines.append([])
            items = [item.strip() for item in line_value.split(",")]
            for item in items:
                bold = "(bold)" in item
                item_name = item.replace("(bold)", "").strip()
                # write_item(item_name, bold=bold)
                all_lines[i].append((item_name, bold))
                all_items[item_name] = 0
    
    scores[abbr]["Lines"] = all_lines
    scores[abbr]["All items"] = all_items

##################################################
# Set up side bar
def clear_my_cache():
    st.cache_data.clear()

with st.sidebar:
    st.markdown("**After editing dropdown options, please reload data using the button below to update within the form.**")
    st.link_button("Edit Dropdown Options", st.secrets['mod3_spreadsheet'])
    st.link_button("Edit Score Options", st.secrets['mod3_scores'])
    st.button('Reload Spreadsheet Data', on_click=clear_my_cache)

    ####################################################
    st.markdown("**Check to include score in the form:** Scores to report:")
    scq_result = st.checkbox("Social Communication Questionnaire (SCQ) - Lifetime Form")
    teacher_srs_eval = st.checkbox("Teacher's SRS Scores")
    informant_vineland_eval = st.checkbox("Informant's Report - Vineland Adaptive Behavior Scales")
    caregiver_vineland_eval = st.checkbox("Caregiver's Vineland Adaptive Behavior Scales")
    teacher_vineland_eval = st.checkbox("Teacher's Vineland Adaptive Behavior Scales")
    wais_check = st.checkbox("Wechsler Adult Intelligence Scale – Fifth Edition (WAIS-5)")
    pai_check = st.checkbox("Personality Assessment Inventory (PAI)")
    
    for item in scores:
        check_scores[item] = st.checkbox(scores[item]["Test name"])

col1,col2 = st.columns(2)
col1.title('Module 4 Report Builder')

def format_date_with_ordinal(date_obj):
    day = date_obj.day
    suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return date_obj.strftime(f"%B {day}{suffix}, %Y")

# Set up dictionary to store data 
data = {}
optional = {}
teacher_score = {}
bullet = {}
lines = {}
comma = {}

# set up recommendation system
check_rec = {}
with open("misc_data/rec_per_module.yaml", "r") as file:
    recommendation_options = yaml.safe_load(file)['mod_3']

####################################################
st.header("Appointment Summary")

data['{{Patient First Name}}'] = st.text_input('Patient First Name')

data['{{Patient Last Name}}'] = st.text_input('Patient Last Name')

preferred = st.selectbox(
    "Patient's Preferred Pronoun",
    ("They/them", "He/him", "She/her"),
)

data['{{Location of the evaluation}}'] = st.radio(
    "Location of the evaluation",
    ['home', 'school', 'the office'],
    index=None,
)

# Audio section 
st.markdown(f"**Behavioral Observation:** Things to mention: eye contact, attention to task, social affect and restricted and repetitive behavior.")
audio_behavior = st.audio_input("Behavioral Observation")
if audio_behavior:
    # 3. Create a download button
    st.download_button(
        label="Download Behavioral Observation Recording",
        key="audio_behavior",
        data=audio_behavior,
        file_name=f"{data['{{Patient First Name}}']} {data['{{Patient Last Name}}']} - Behavioral Observation.wav",
        mime="audio/wav",
    )

st.markdown(f"**Developmental History:** Things to mention: social communication skills, repetitive behavior and related behavioral concerns.")
audio_development = st.audio_input("Developmental History")
if audio_development:
    # 3. Create a download button
    st.download_button(
        label="Download Developmental History Recording",
        key="audio_development",
        data=audio_development,
        file_name=f"{data['{{Patient First Name}}']} {data['{{Patient Last Name}}']} - Developmental History.wav",
        mime="audio/wav",
    )

if st.button("Transcribe"):
    if audio_behavior and audio_development:
        transcript_behavior = transcribe_audio(audio_behavior, name='behavior')
        st.markdown(f"**Transcription:** {transcript_behavior}")

        transcript_development = transcribe_audio(audio_development, name='development')
        st.markdown(f"**Transcription:** {transcript_development}")
        
        response = client.responses.create(
            prompt={
                "id": st.secrets["behavior_prompt_mod3_id"],
                # "version": "3",
                "variables": {
                    "first_name": data['{{Patient First Name}}'],
                    "pronouns": preferred,
                    "evaluation_location": data['{{Location of the evaluation}}'],
                    "transcription": transcript_behavior
                }
            }
        )
        st.session_state.behavior_observation_mod3 = response.output_text

        response = client.responses.create(
            prompt={
                "id": st.secrets["development_prompt_mod3_id"],
                # "version": "5",
                "variables": {
                    "first_name": data['{{Patient First Name}}'],
                    "pronouns": preferred,
                    "transcription": transcript_development
                }
            }
        )
        st.session_state.development_history_mod3 = response.output_text
   
####################################################
with st.form('BasicInfo'):
    st.header("Patient's data")

    data["{{Patient Age}}"] = st.number_input("Patient's Age", 0, 100)
    data['{{Patient age unit}}'] = st.radio(
        "Year/month?",
        ("year", "month")
    )

    data['{{Caregiver type}}'] = st.selectbox(
        "Patient's Caregiver",
        ("mother", "father", "parent", "grandparent", "legal custodian", "foster parent"),
        placeholder="Select from the choices or enter a new one",
        index=None,
        accept_new_options=True,
    )

    bullet['CaregiverPrimaryConcerns'] = st.multiselect(
        "Caregiver\'s Primary Concerns",
        dropdowns["Caregiver\'s Primary Concerns"],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )
    
    data['{{Residence City/State}}'] = st.text_input("Residence City/State")

    data['{{Narrative}}'] = st.text_area('Narrative to finish \"Patient lives with...\"')

    ##########################################################
    st.header("BRH Evaluation Details")

    data['{{Evaluation Date}}'] = format_date_with_ordinal(st.date_input("Evaluation Date"))

    data['{{Results Shared Date}}'] = format_date_with_ordinal(st.date_input("Results Shared Date"))
    
    data['{{Date Report Sent to Patient}}'] = format_date_with_ordinal(st.date_input("Date Report Sent to Patient"))

    lines["{{Result of the evaluation}}"] = st.multiselect(
        "Result of the evaluation",
        dropdowns['Result of the evaluation'],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )
    
    ##########################################################
    if scq_result:
        st.header("(SCQ) - Lifetime Form")
        data["Results (SCQ) Lifetime Form"] = st.text_input(
            "Results (SCQ) - Lifetime Form"
        )

    ##########################################################
    st.header("SRS-2 Caregiver Score")

    data["{{SRS-2 Score Caregiver}}"] = st.text_input("Caregiver's SRS-2 Score")
    
    data["{{Social Communication and Interaction Score Caregiver}}"] = st.text_input("Social Communication and Interaction Score Caregiver")
    
    data["{{Restricted Interests and Repetitive Behavior Score Caregiver}}"] = st.text_input("Restricted Interests and Repetitive Behavior Score Caregiver")

    data["{{Caregiver's level of concern}}"] = st.radio(
        "Caregiver's level of concern",
        ['no', 'mild', 'moderate', 'severe']
    )

    data["{{Evaluator's level of concern}}"] = st.radio(
        "Evaluator's level of concern",
        ['no', 'mild', 'moderate', 'severe']
    )

    ##########################################################
    if teacher_srs_eval:
        st.header("Teacher SRS Score")
        st.markdown("*Skip this section if teacher did not give SRS Score*")

        teacher_score['{{SRS-2 Score Teacher}}'] = st.text_input("Teacher's SRS-2 Score")

        teacher_score['{{Social Communication and Interaction Score Teacher}}'] = st.text_input("Social Communication and Interaction Score Teacher")

        teacher_score['{{Restricted Interests and Repetitive Behavior Score Teacher}}'] = st.text_input("Restricted Interests and Repetitive Behavior Score Teacher")

        teacher_score["{{Teacher level of concern}}"] = st.radio(
            "Teacher's level of concern",
            ['no', 'mild', 'moderate', 'severe']
        )
    
    ##########################################################
    if caregiver_vineland_eval:
        st.header("Vineland Adaptive Behavior Scales - Caregiver's Scores")

        data["{{Vineland Score Caregiver}}"] = st.text_input("Vineland Score Caregiver")
        
        data["{{Communication Score Caregiver}}"] = st.text_input("Communication Score Caregiver")
        
        data["{{Daily Living Skills Score Caregiver}}"] = st.text_input("Daily Living Skills Score Caregiver")
        
        data["{{Socialization Score Caregiver}}"] = st.text_input("Socialization Score Caregiver")

        if teacher_vineland_eval:
            st.header("Vineland Adaptive Behavior Scales - Teacher's Scores")

            data["{{Vineland Score Teacher}}"] = st.text_input("Vineland Score Teacher")
            
            data["{{Communication Score Teacher}}"] = st.text_input("Communication Score Teacher")
            
            data["{{Daily Living Skills Score Teacher}}"] = st.text_input("Daily Living Skills Score Teacher")
            
            data["{{Socialization Score Teacher}}"] = st.text_input("Socialization Score Teacher")

    ######################################################
    st.header("Medical/Developmental History")
    
    data['{{Developmental History}}'] = st.text_input(
        "Developmental History")
    
    data['{{Educational History}}'] = st.text_input(
        "Educational History")

    data['{{Diagnostic History}}'] = st.text_input(
        "Diagnostic History")
    
    data['{{Medical History}}'] = st.text_input(
        "Relevant Medical History")

    lines['{{Medications}}'] = st.multiselect(
        "Medications",
        ['None noted or reported.'],
        placeholder="Can input multiple options",
        accept_new_options=True
    )

    ###############################################
    st.header("Educational Background")

    data['{{School District}}'] = st.selectbox(
        "School District",
        ['Rochester City'],
        index=None,
        placeholder="Select a school district or enter a new one",
        accept_new_options=True,
    )

    # data['{{School Name}}'] = st.text_input("School Name")

    if teacher_srs_eval or teacher_vineland_eval:
        teacher_score['{{Teacher name, title}}'] = st.text_input("Teacher name, title")

    data['{{Grade}}'] = st.text_input(
        "Grade",
    )

    data['School Year'] = st.text_input(
        "School Year",
    )

    data['{{Education Setting}}'] = st.selectbox(
        "Education Setting",
        ["General Education", "Integrated Co-Taught", "12:1:1", "8:1:1", "6:1:1"],
        index=None,
        placeholder="Select a grade or enter a new one",
        accept_new_options=True,
    )

    comma['{{Services}}'] = st.multiselect(
        "Services",
        dropdowns['Services'],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )

    data['{{Classification}}'] = st.selectbox(
        "Classification",
        dropdowns['Classification'],
        index=None,
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )

    ################################################# 
    if wais_check:
        st.header("WAIS-5 Score Report")

        wais_data = {}
        wais_data['overall'] = st.text_area(
            "WAIS-5 Overall Score - Input percentile without any suffix. For example: Percentile: 42", 
            """Full Scale IQ:
  Standard Score: 
  Confidence Interval: 
  Percentile: 

Verbal Comprehension:
  Standard Score: 
  Confidence Interval: 
  Percentile: 

Visual Spatial:
  Standard Score: 
  Confidence Interval: 
  Percentile:

Fluid Reasoning:
  Standard Score: 
  Confidence Interval:
  Percentile: 

Working Memory:
  Standard Score: 
  Confidence Interval: 
  Percentile: 

Processing Speed:
  Standard Score: 
  Confidence Interval: 
  Percentile: 
""",
            height=800,
        )

        wais_data['subtest'] = st.text_area(
            "WAIS-5 Subtest Score", 
            """Verbal Comprehension:
  Similarities: 
  Vocabulary: 

Working Memory:
  Digit Sequencing: 
  Running Digits: 

Visual Spatial:
  Block Design: 
  Visual Puzzles: 

Processing Speed:
  Symbol Search: 
  Coding: 

Fluid Reasoning:
  Matrix Reasoning: 
  Figure Weights: 
""",
        height=500,
    )
        
    #############################################
    # First table
    if informant_vineland_eval:
        st.header("Informant's Report - Vineland Adaptive Behavior Scales")

        # Load data
        df = pd.read_csv("misc_data/vineland_informant.csv")

        # JavaScript code to apply bold styling if "bold" column is True
        row_style_jscode = JsCode("""
        function(params) {
            if (params.data.bold === true) {
                return {
                    'font-weight': 'bold',
                    'font-size': 16,
                }
            } else {
                return {
                    'font-size': 16,
                }
            }
            return {};
        }
        """)

        # Build Grid Options
        gb = GridOptionsBuilder.from_dataframe(df)
        gb.configure_grid_options(getRowStyle=row_style_jscode)
        gb.configure_column("data", editable=True)
        gb.configure_column("field", width=300) 
        gridOptions = gb.build()
        st.write(gridOptions["columnDefs"])

        # Display grid
        # with st_normal():
        grid_return = AgGrid(
            df,
            gridOptions=gridOptions,
            editable=True,
            height=800,
            theme="balham",
            # custom_css=custom_css,
            allow_unsafe_jscode=True
        )

    ##########################################################
    # Score section
    for test in check_scores:
        if check_scores[test]:
            st.header(scores[test]["Test name"])
            st.markdown(f"*Skip this section if there is no {test} Score*")

            scores[test]["Test Date"] = st.date_input(f"{test} Test Date").strftime("%m/%Y")
            
            for item in scores[test]["All items"]:
                scores[test]["All items"][item] = st.text_input(item)    

    ########################################################
    st.header("Behavioral Presentation")
    data['behavior_observation'] = st.text_area(
        "Behavioral Observation: Edit the response before submitting the form", 
        # behavior_observation,
        st.session_state.behavior_observation_mod3,
        height=350,
    )

    ########################################################
    st.header("Developmental History")
    data['development_history'] = st.text_area(
        "Developmental History: Edit the response before submitting the form", 
        # development_history,
        st.session_state.development_history_mod3,
        height=350,
    )

    ########################################################
    st.header("Diagnostic Formulation")
    diagnostic_formulation = st.selectbox(
        "Diagnostic Formulation",
        ["Cognitive Deficit", "Adaptive Functioning", "Both (Cognitive and Adaptive Functioning)"],
        index=None,
        placeholder="Select options from the list",
    )

    ############################################
    st.header("DSM Criteria")

    # Add DSM questions
    with open("misc_data/dsm.yaml", "r") as file:
        dsm_title = yaml.safe_load(file)
    
    for key, value in dsm_title.items():
        bullet[key] = st.multiselect(    
            value,
            dropdowns[key],
            placeholder="Select multiple options from the list or enter a new one",
            accept_new_options=True
        )

    comma['{{Symptoms present in the early developmental period}}'] = st.multiselect(
        "Symptoms present in the early developmental period",
        [
            "Confirmed by record review",
            "None",
        ],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )

    comma['{{Symptoms cause clinically significant impairment}}'] = st.multiselect(
        "Symptoms cause clinically significant impairment",
        [
            "Confirmed by record review",
            "None",
        ],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )

    ##########################################################
    st.header("Recommendations")

    for key, label in recommendation_options.items():
        check_rec[key] = st.checkbox(label)

    # data['{{}}'] = st.text_input("")
    # data['{{}}'] = st.text_input("")
    # data['{{}}'] = st.text_input("")

    submit = st.form_submit_button('Submit')

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def add_behavior_presentation(paragraph, transcript):
    if transcript != "":    
        # separate transcript
        small_para = transcript.split('\n\n')

        st.write(small_para)

        paragraph.insert_paragraph_before().add_run(small_para[0], style='CustomStyle')
        paragraph.insert_paragraph_before()

        for sub_para in small_para[1:]:
            sub_para = sub_para.split(":")
            p = paragraph.insert_paragraph_before()
            p.add_run(sub_para[0], style='CustomStyle').italic = True
            p.add_run(f":{sub_para[1]}\n", style='CustomStyle')
        
    delete_paragraph(paragraph)

def add_developmental_history(paragraph, transcript):
    if transcript != "":    
        # separate transcript
        small_para = transcript.split('\n\n')
        st.write(small_para)

        for sub_para in small_para:
            sub_para = sub_para.split(":")
            p = paragraph.insert_paragraph_before()
            p.add_run(sub_para[0], style='CustomStyle').italic = True
            p.add_run(f":{sub_para[1]}\n", style='CustomStyle')
        
    delete_paragraph(paragraph)

def add_school(paragraph):
    p = paragraph.insert_paragraph_before()
    tab_stops = p.paragraph_format.tab_stops
    # tab_stops.clear()  # Start fresh for this paragraph only
    tab_stops.add_tab_stop(Inches(2.5))
    # Add data
    p.add_run("District", style='CustomStyle').font.underline = True
    p.add_run(f": {data['{{School District}}']}\t", style='CustomStyle')
    p.add_run("Classification", style='CustomStyle').font.underline = True
    p.add_run(": {{Classification}}\n\n", style='CustomStyle')

    p.add_run("Setting", style='CustomStyle').font.underline = True
    p.add_run(f": {data['{{Education Setting}}']}\t", style='CustomStyle')
    p.add_run("Grade", style='CustomStyle').font.underline = True
    p.add_run(f": {data['{{Grade}}']} (", style='CustomStyle')
    p.add_run(f"{data['School Year']} school year", style='CustomStyle').italic = True
    p.add_run(")", style='CustomStyle')
    delete_paragraph(paragraph)

def add_scq_form(paragraph):
    r = paragraph.insert_paragraph_before().add_run('Social Communication Questionnaire (SCQ) – Lifetime Form', style='CustomStyle')
    r.italic = True
    r.font.underline = True
    p = paragraph.insert_paragraph_before()
    p.add_run("The SCQ evaluates for symptoms of autism spectrum disorder across developmental history. Scores above 15 are suggestive of an autism diagnosis. Based on the {{Caregiver type}}’s report, ", style='CustomStyle')
    p.add_run("{{Patient First Name}}’s score was ", style='CustomStyle').bold = True
    p.add_run(f"{data['Results (SCQ) Lifetime Form']}. ", style='CustomStyle').bold = True
    r = p.add_run("This score is clearly consistent with autism at present.\n", style='CustomStyle')
    r.bold = True
    r.italic = True
    delete_paragraph(paragraph)

def add_vineland_no_teacher(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run("Vineland Adaptive Behavior Scales – 3rd Ed. (VABS-3) – Parent's Report", style='CustomStyle')
    r.font.underline = True
    r.italic = True

    paragraph.insert_paragraph_before().add_run('The VABS-3 yields information about an individual’s adaptive functioning, which is the ability to independently perform daily activities for personal and social sufficiency. The Adaptive Behavior Composite measures overall adaptive functioning, while separate scores provide more details about communication, daily living skills, and socialization. \n\nStandard scores on the VABS-3 have a mean of 100 and a standard deviation of 15.  Scores between 85 and 115 are within the average range for this test, scores between 70 and 84 are considered moderately low, and scores below 70 are considered very low.\n\n', style='CustomStyle')
    
    paragraph.insert_paragraph_before().add_run('\tAdaptive Behavior Composite: {{Vineland Score Caregiver}} ({{Caregiver type}})', style='CustomStyle').bold = True
    paragraph.insert_paragraph_before().add_run('\t\tCommunication: {{Communication Score Caregiver}} ({{Caregiver type}})', style='CustomStyle')
    paragraph.insert_paragraph_before().add_run('\t\tDaily Living Skills: {{Daily Living Skills Score Caregiver}} ({{Caregiver type}})', style='CustomStyle')
    paragraph.insert_paragraph_before().add_run('\t\tSocialization: {{Socialization Score Caregiver}} ({{Caregiver type}})', style='CustomStyle')
    delete_paragraph(paragraph)

def add_vineland_yes_teacher(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run("Vineland Adaptive Behavior Scales – 3rd Ed. (VABS-3) – Parent's Report & Teacher Report", style='CustomStyle')
    r.font.underline = True
    r.italic = True

    paragraph.insert_paragraph_before().add_run('The VABS-3 yields information about an individual’s adaptive functioning, which is the ability to independently perform daily activities for personal and social sufficiency. The Adaptive Behavior Composite measures overall adaptive functioning, while separate scores provide more details about communication, daily living skills, and socialization. \n\nStandard scores on the VABS-3 have a mean of 100 and a standard deviation of 15.  Scores between 85 and 115 are within the average range for this test, scores between 70 and 84 are considered moderately low, and scores below 70 are considered very low.\n\n', style='CustomStyle')
    
    paragraph.insert_paragraph_before().add_run('\tAdaptive Behavior Composite: {{Vineland Score Caregiver}} ({{Caregiver type}}), {{Vineland Score Teacher}} (teacher)', style='CustomStyle').bold = True
    paragraph.insert_paragraph_before().add_run('\t\tCommunication: {{Communication Score Caregiver}} ({{Caregiver type}}), {{Communication Score Teacher}} (teacher)', style='CustomStyle')
    paragraph.insert_paragraph_before().add_run('\t\tDaily Living Skills: {{Daily Living Skills Score Caregiver}} ({{Caregiver type}}), {{Daily Living Skills Score Teacher}} (teacher)', style='CustomStyle')
    paragraph.insert_paragraph_before().add_run('\t\tSocialization: {{Socialization Score Caregiver}} ({{Caregiver type}}), {{Socialization Score Teacher}} (teacher)', style='CustomStyle')
    delete_paragraph(paragraph)

def add_srs_no_teacher(paragraph):
    r = paragraph.insert_paragraph_before().add_run('Social Responsiveness Scale – Second Edition (SRS-2) – Parent Report', style='CustomStyle')
    r.italic = True
    r.font.underline = True
    p = paragraph.insert_paragraph_before()
    p.add_run('The SRS-2 is an objective measure that identifies social impairments associated with autism spectrum disorder and quantifies ASD-related severity throughout the lifespan.\n\nStudies show that the SRS-2 discriminates both ', style='CustomStyle')
    p.add_run("within ", style='CustomStyle').italic = True
    p.add_run("the autism spectrum and between ASD and other disorder, which makes the test useful for differential diagnosis. Raters evaluate symptoms using a scale representing a range of severity. Although not used for diagnosis, subscale scores are helpful in designing and evaluating treatment.\n\nSRS-2 scores are reported here as T-scores with a mean of 50 and a standard deviation of 10 with higher scores indicating greater levels of concern for how social behavior impacts or interferes with everyday interactions. The following interpretative guidelines are offered here for the benefit of the reader: Less than 59 indicates within normal limits, between 60 and 65 as mild concern, between 65 and 75 as moderate concern, and greater than 76 as severe. {{Patient First Name}}’s {{Caregiver type}} and teacher reported the following:", style='CustomStyle')
    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before()
    r = p.add_run('\tSRS-2 Total Score', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run(': {{SRS-2 Score Caregiver}} ({{Caregiver type}})', style='CustomStyle').bold = True
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run('\tSocial Communication and Interaction: {{Social Communication and Interaction Score Caregiver}} ({{Caregiver type}})', style='CustomStyle')
    paragraph.insert_paragraph_before().add_run('\tRestricted Interests and Repetitive Behavior: {{Restricted Interests and Repetitive Behavior Score Caregiver}} ({{Caregiver type}})', style='CustomStyle')
    paragraph.insert_paragraph_before()
    observe = paragraph.insert_paragraph_before()
    observe.add_run("Based on the report provided by {{Preferred Pronouns 2}} {{Caregiver type}}, ", style='CustomStyle')
    observe.add_run("{{Patient First Name}}’s social communication and related behaviors indicated {{Caregiver's level of concern}} concerns. ", style='CustomStyle').italic = True
    observe.add_run("My observation aligned with a {{Evaluator's level of concern}} concern.\n", style='CustomStyle').bold = True
    delete_paragraph(paragraph)

def add_srs_yes_teacher(paragraph, score_data):
    r = paragraph.insert_paragraph_before().add_run('Social Responsiveness Scale – Second Edition (SRS-2) – Parent Report', style='CustomStyle')
    r.italic = True
    r.font.underline = True
    p = paragraph.insert_paragraph_before()
    p.add_run('The SRS-2 is an objective measure that identifies social impairments associated with autism spectrum disorder and quantifies ASD-related severity throughout the lifespan.\n\nStudies show that the SRS-2 discriminates both ', style='CustomStyle')
    p.add_run("within ", style='CustomStyle').italic = True
    p.add_run("the autism spectrum and between ASD and other disorder, which makes the test useful for differential diagnosis. Raters evaluate symptoms using a scale representing a range of severity. Although not used for diagnosis, subscale scores are helpful in designing and evaluating treatment.\n\nSRS-2 scores are reported here as T-scores with a mean of 50 and a standard deviation of 10 with higher scores indicating greater levels of concern for how social behavior impacts or interferes with everyday interactions. The following interpretative guidelines are offered here for the benefit of the reader: Less than 59 indicates within normal limits, between 60 and 65 as mild concern, between 65 and 75 as moderate concern, and greater than 76 as severe. {{Patient First Name}}’s {{Caregiver type}} and teacher reported the following:", style='CustomStyle')
    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before()
    p.add_run('\tSRS-2 Total Score: {{SRS-2 Score Caregiver}} ({{Caregiver type}}), ', style='CustomStyle').bold = True
    p.add_run(f"{score_data['{{SRS-2 Score Teacher}}']} (teacher)", style='CustomStyle').bold = True
    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before()
    p.add_run('\tSocial Communication and Interaction: {{Social Communication and Interaction Score Caregiver}} ({{Caregiver type}}), ', style='CustomStyle')
    p.add_run(f"{score_data['{{Social Communication and Interaction Score Teacher}}']} (teacher)", style='CustomStyle')
    p = paragraph.insert_paragraph_before()
    p.add_run('\tRestricted Interests and Repetitive Behavior: {{Restricted Interests and Repetitive Behavior Score Caregiver}} ({{Caregiver type}}), ', style='CustomStyle')
    p.add_run(f'{score_data["{{Restricted Interests and Repetitive Behavior Score Teacher}}"]} (teacher)', style='CustomStyle')
    paragraph.insert_paragraph_before()
    observe = paragraph.insert_paragraph_before()
    observe.add_run("Based on the report provided by {{Preferred Pronouns 2}} {{Caregiver type}}, ", style='CustomStyle')
    observe.add_run("{{Patient First Name}}’s social communication and related behaviors indicated {{Caregiver's level of concern}} concerns. ", style='CustomStyle').italic = True
    observe.add_run("{{Patient First Name}}’s teacher reported a ", style='CustomStyle')
    observe.add_run(f"{score_data['{{Teacher level of concern}}']} concern, and ", style='CustomStyle')
    observe.add_run("my observation aligned with a {{Evaluator's level of concern}} concern.\n", style='CustomStyle').bold = True
    delete_paragraph(paragraph)

def add_diagnostic_formulation(paragraph, formulation):
    if formulation == "Cognitive Deficit":
        p = paragraph.insert_paragraph_before()
        p.add_run('{{Patient First Name}} has the greatest difficulty with skills and behaviors that fall within the domain of cognitive functioning. {{Patient First Name}}’s score on the Full-Scale IQ and Wechsler Adult Intelligence Scale, Fifth Edition is greater than 2.0 standard deviations below the normed average. I believe that {{Patient First Name}}’s handicap with cognitive functioning is best explained by the presence of {{Preferred Pronouns 2}} meeting the criteria for autism spectrum disorder.', style='CustomStyle')

    elif formulation == "Adaptive Functioning":
        p = paragraph.insert_paragraph_before()
        p.add_run('{{Patient First Name}} has the greatest difficulty with skills and behaviors that fall within the domain of adaptive functioning. {{Patient First Name}}’s score on the Adaptive Behavior Compositive of the Vineland Adaptive Behavior Scales, Third Edition is greater than 2.0 standard deviations below the normed average. I believe that {{Patient First Name}}’s handicap with adaptive functioning is best explained by the presence of {{Preferred Pronouns 2}} meeting the criteria for autism spectrum disorder.', style='CustomStyle')

    else:
        p = paragraph.insert_paragraph_before()
        p.add_run('{{Patient First Name}} has the greatest difficulty with skills and behaviors that fall within the domain of cognitive and adaptive functioning. {{Patient First Name}}’s score on the Full-Scale IQ and Wechsler Adult Intelligence Scale, Fifth Edition and Adaptive Behavior Compositive of the Vineland Adaptive Behavior Scales, Third Edition is greater than 2.0 standard deviations below the normed average. I believe that {{Patient First Name}}’s handicap with cognitive and adaptive functioning is best explained by the presence of {{Preferred Pronouns 2}} meeting the criteria for autism spectrum disorder.', style='CustomStyle')
    delete_paragraph(paragraph)

def add_score(paragraph, score_data):
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run(f'\t({score_data["Test Date"]}) \u2013 {score_data["Test name"]}', style='CustomStyle').italic = True
    
    # Go over each line 
    for line in score_data["Lines"]:
        # get a new paragraph and indent it 
        p = paragraph.insert_paragraph_before()
        p.paragraph_format.left_indent = Inches(0.5)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(3.5))

        # add each score
        for item_tuple in line:
            item = item_tuple[0]
            p.add_run(f'{item}: {score_data["All items"][item]}\t', style='CustomStyle').bold = item_tuple[1]


if submit:
    # Update session state 
    st.session_state.behavior_observation_mod3 = data['behavior_observation']
    st.session_state.development_history_mod3 = data['development_history'] 

    # handle word to replace 
    # pronouns
    with open("misc_data/pronouns.yaml", "r") as file:
        pronoun = yaml.safe_load(file)

    replace_word = {
        "{{Preferred Pronouns 1}}": pronoun[preferred]['pronoun1'],
        "{{Preferred Pronouns 1 CAP}}": pronoun[preferred]['pronoun1cap'],
        "{{Preferred Pronouns 2}}": pronoun[preferred]['pronoun2'],
        "{{Preferred Pronouns 2 CAP}}": pronoun[preferred]['pronoun2cap'],
        "{{Gender}}": pronoun[preferred]['gender'],
    }

    replace_word.update(data)
    
    replace_percent = {}

    # Get WAIS 
    if wais_check:
        wais_analysis = ""
        if wais_data["subtest"] and wais_data['overall']:
            response = client.responses.create(
                prompt={
                    "id": st.secrets["wais_analysis_id"],
                    "variables": {
                        "first_name": data['{{Patient First Name}}'],
                        "pronouns": preferred,
                        "wais_subtest": wais_data['subtest'],
                        "wais_overall": wais_data['overall'],
                    }
                }
            )
            wais_analysis = response.output_text

        # Get Table data
        wais_subtest_score = yaml.safe_load(wais_data['subtest'])
        wais_score = yaml.safe_load(wais_data['overall'])

        info_list = ['IQ', 'Verbal Comp', 'Visual Spatial']

        for info in wais_score:
            # Info meaning type of score, i.e. "Full Scale IQ", "Verbal Comprehension Index, etc"
            replace_word[f"[[{info} Standard]]"] = str(wais_score[info]['Standard Score'])
            replace_word[f"[[{info} CI]]"] = str(wais_score[info]['Confidence Interval'])
            replace_percent[f"[[{info} Percent]]"] = str(wais_score[info]['Percentile'])
        
        for info in wais_subtest_score:
            for subtest in wais_subtest_score[info]:
                replace_word[f"[[{subtest}]]"] = str(wais_subtest_score[info][subtest])


    # Get Vineland Informant
    if informant_vineland_eval:
        # Vineland Informant Score
        vineland_info_dict = {
            "[[{}]]".format(row["field"]): row["data"] 
            for _, row in grid_return['data'].iterrows()
        }

        vineland_perc_dict = {
            "[[{}]]".format(row["field"]): row["data"] 
            for _, row in grid_return['data'].iterrows()
            if "Percentile" in row["field"]
        }

        replace_percent.update(vineland_perc_dict)
        replace_word.update(vineland_info_dict)

    # Display data 
    # yaml_string = yaml.dump(replace_word, sort_keys=False)
    # yaml_string = yaml_string + '\n' + yaml.dump(scores, sort_keys=False)
    # yaml_string = yaml_string + '\n' + yaml.dump(bullet, sort_keys=False)
    # if wais_check:
    #     yaml_string += "\n" + wais_analysis
    # yaml_data = st.code(yaml_string, language=None)
    

    #### Edit document 
    # if wais_check:
    #     doc = Document('templates/template_mod_4.docx')
    # else:
    #     doc = Document('templates/template_mod_4_no_wais.docx')
    
    doc = Document('templates/template_mod_4.docx')
    if doc:
        # Get file name
        today_date = format_date_with_ordinal(datetime.date.today())
        filename = f"{data['{{Patient First Name}}']} {data['{{Patient Last Name}}']} {today_date}.docx"
        
        ### create document 
        norm_style = doc.styles['Normal']
        norm_style.paragraph_format.line_spacing = 1

        custom_style = doc.styles.add_style('CustomStyle', WD_STYLE_TYPE.CHARACTER)
        custom_style.font.size = Pt(12)
        custom_style.font.name = 'Georgia'

        custom_style_2 = doc.styles.add_style('CustomStyle2', WD_STYLE_TYPE.CHARACTER)
        custom_style_2.font.size = Pt(11.5)
        custom_style_2.font.name = 'Georgia'

        list_style = doc.styles['Bullet New']
        list_style.paragraph_format.line_spacing = 1

        if not informant_vineland_eval:
            # if we don't have vineland score
            # Remove tables 
            allTables = doc.tables
            for activeTable in allTables:
                if activeTable.cell(0,0).paragraphs[0].text == 'Adaptive Behavior Composite':
                    activeTable._element.getparent().remove(activeTable._element)

            # Test paragraph 
            for i, paragraph in enumerate(doc.paragraphs):
                # Vineland Informant Report 
                if "[[Vineland_Start]]" in paragraph.text:
                    vineland_start = i 
                    break

            if vineland_start:
                print(vineland_start)
                for index in range(vineland_start, vineland_start+4, 1):
                    try:
                        delete_paragraph(doc.paragraphs[index])
                    except:
                        print("Out of range at index", i)
            else:
                print("Cannot find Vineland Start")

            for i, paragraph in enumerate(doc.paragraphs):
                if "The VABS-3 yields information about an individual’s adaptive functioning" in paragraph.text:
                    delete_paragraph(paragraph)
                if "[[Vineland Analysis]]" in paragraph.text:
                    delete_paragraph(paragraph)

        if wais_check or informant_vineland_eval:
            # Replace percent in table 
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Loop through all percentage needed to be replaced in the table
                            for key in replace_percent:
                                if key in paragraph.text:
                                    p = paragraph.insert_paragraph_before()
                                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    p.add_run(replace_percent[key], style='CustomStyle')
                                    suffix = get_ordinal(replace_percent[key])
                                    p.add_run(suffix, style='CustomStyle').font.superscript = True
                                    delete_paragraph(paragraph)

        # Add scores 
        for paragraph in doc.paragraphs:
            if "Scores are reported here as standard scores" in paragraph.text:
                total = 0
                for test in check_scores:
                    if check_scores[test]:
                        total += 1 
                        if total == 1:
                            paragraph.insert_paragraph_before().add_run("\nPsychoeducational Testing:", style='CustomStyle').font.underline = True
                        add_score(paragraph, score_data=scores[test])

                if total == 0:
                    delete_paragraph(paragraph)

            if "[[Assessment Procedures]]" in paragraph.text:
                if informant_vineland_eval:
                    replace_ordinal_with_superscript(paragraph, "Vineland Adaptive Behavior Assessment Scale – 3rd Ed: Completed by {{Preferred Pronouns 2}}  {{Caregiver type}}")
                if pai_check:
                    replace_ordinal_with_superscript(paragraph, "Personality Assessment Inventory (PAI): Completed by self")
                delete_paragraph(paragraph)
        
            if "[[Behavioral Presentation]]" in paragraph.text:
                add_behavior_presentation(paragraph, st.session_state.behavior_observation_mod3)
            
            if "[[Developmental History]]" in paragraph.text:
                add_developmental_history(paragraph, st.session_state.development_history_mod3)

            if "[[Diagnostic Formulation]]" in paragraph.text:
                add_diagnostic_formulation(paragraph, diagnostic_formulation)

            if "[[Recommendations]]" in paragraph.text:
                for rec, checked in check_rec.items():
                    if checked:
                        func = globals().get(f"add_{rec}")
                        if callable(func):
                            func(paragraph)
                
                delete_paragraph(paragraph)

            if "[[WAIS-Analysis]]" in paragraph.text:
                if wais_check:
                    replace_ordinal_with_superscript(paragraph, wais_analysis)

            if "[[SCQ Report Information]]" in paragraph.text:
                # Add SCQ
                if scq_result:
                    add_scq_form(paragraph)
                else:
                    delete_paragraph(paragraph)
                
            if "[[SRS Report Information]]" in paragraph.text:
                # Add SRS
                if teacher_srs_eval:
                    add_srs_yes_teacher(paragraph, teacher_score)
                else:
                    add_srs_no_teacher(paragraph)
            
            if "Social Responsiveness Scale" in paragraph.text:
                if teacher_srs_eval:
                    paragraph.add_run(" & teacher", style='CustomStyle')
            
            if "Developmental History & Review of Records" in paragraph.text:
                if teacher_srs_eval:
                    paragraph.add_run(f"\nSchool Report on SRS-2 provided by {teacher_score['{{Teacher name, title}}']}", style='CustomStyle')
                if teacher_vineland_eval:
                    paragraph.add_run(f"\nReport on Vineland Adaptive Behavior Scale provided by {teacher_score['{{Teacher name, title}}']}", style='CustomStyle')

            if "[[District Grade School Setting]]" in paragraph.text:
                add_school(paragraph)

            if "[[Vineland Score Breakdown]]" in paragraph.text:
                if teacher_vineland_eval:
                    add_vineland_yes_teacher(paragraph)
                else:
                    add_vineland_no_teacher(paragraph)

            if "[[Vineland Analysis]]" in paragraph.text:
                if informant_vineland_eval:
                    # Add page break
                    paragraph.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)

                    p = paragraph.insert_paragraph_before()
                    r = p.add_run("Interpretation of VABS-3 Results – Informant Report", style='CustomStyle')
                    r.bold = True
                    r.italic = True

                    p = paragraph.insert_paragraph_before()
                    p.add_run(f"\nAdaptive Behavior Composite: {replace_word['[[Adaptive Behavior Composite]]']} (", style='CustomStyle')
                    write_ordinal_in_paragraph(p, replace_word['[[Adaptive Behavior Composite Percentile]]'])
                    p.add_run(" percentile)", style='CustomStyle')

                    p = paragraph.insert_paragraph_before()
                    p.add_run(f"\n\tCommunication: {replace_word['[[Communication]]']} (", style='CustomStyle')
                    write_ordinal_in_paragraph(p, replace_word['[[Communication Percentile]]'])
                    p.add_run(" percentile)", style='CustomStyle')

                    p = paragraph.insert_paragraph_before()
                    p.add_run(f"\tDaily Living Skills: {replace_word['[[Daily Living Skills]]']} (", style='CustomStyle')
                    write_ordinal_in_paragraph(p, replace_word['[[Daily Living Skills Percentile]]'])
                    p.add_run(" percentile)", style='CustomStyle')

                    p = paragraph.insert_paragraph_before()
                    p.add_run(f"\tSocialization: {replace_word['[[Socialization]]']} (", style='CustomStyle')
                    write_ordinal_in_paragraph(p, replace_word['[[Socialization Percentile]]'])
                    p.add_run(" percentile)", style='CustomStyle')

                    paragraph.insert_paragraph_before()

                # # if not, delete this paragraph 
                # else:
                #     delete_paragraph(paragraph)




                    
                    

        # Edit document
        for word in replace_word:
            docxedit.replace_string(doc, old_string=word, new_string=replace_word[word])

        # Replace for lists separated by comma:
        for word in comma:
            new_word = ", ".join(comma[word])
            docxedit.replace_string(doc, old_string=word, new_string=new_word)

        # Replace for lists separated by new line:
        for word in lines:
            new_word = "\n".join(lines[word])
            docxedit.replace_string(doc, old_string=word, new_string=new_word)

        # Save content to file
        doc.save(filename)

        # Replace for lists separated by bullet points
        tpl=DocxTemplate(filename)
        print("Load template!")

        tpl.render(bullet)
        print("Bullet rendered!")

        tpl.save(filename)
        print("File saved at", filename)

        # Download 
        bio = io.BytesIO()
        document = Document(filename)
        document.save(bio)
        
        st.download_button(
            label="Click here to download",
            key="report_download",
            data=bio.getvalue(),
            file_name=filename,
            mime="docx"
        )