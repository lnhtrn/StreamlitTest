import streamlit as st
import whisper
from openai import OpenAI
import yaml
from docx import Document
import io
import docx
import yaml
import io
import docxedit
import datetime
from docx.enum.style import WD_STYLE_TYPE
from streamlit_gsheets import GSheetsConnection
from docxtpl import DocxTemplate
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement, qn

##########################################################
st.set_page_config(
    page_title="Audio Testing with Whisper",
    page_icon="🎧",
    layout="centered",
    initial_sidebar_state="expanded",
)

##########################################################
# Session state keys: 
if 'behavior_observation' not in st.session_state:
    st.session_state.behavior_observation = ""
if 'development_history' not in st.session_state:
    st.session_state.development_history = ""
if 'final_text' not in st.session_state:
    st.session_state.final_text = ""

data = {}

# Load OpenAI client 
client = OpenAI(api_key=st.secrets["openai_key"])

##################################################################
def transcribe_audio(audio_file, name='temp'):
    if audio_file:
        # Transcribe
        with st.spinner("Transcribing...", show_time=True):
            # result = whisper_model.transcribe(f"{name}.wav")
            result = client.audio.transcriptions.create(
                model="whisper-1", 
                file=audio_file, 
                response_format="text"
            )
        return result 

##################################################################
# Form Builder Simulation 
st.header("Appointment Summary")
data['{{Patient First Name}}'] = st.text_input('Patient First Name')
data['{{Patient Last Name}}'] = st.text_input('Patient Last Name')
preferred = st.selectbox(
        "Patient's Preferred Pronoun",
        ("They/them", "He/him", "She/her"),
    )
audio_behavior = st.audio_input("Behavioral Observation")
audio_development = st.audio_input("Developmental History")

if st.button("Transcribe"):
    if audio_behavior and audio_development:
        transcript_behavior = transcribe_audio(audio_behavior, name='behavior')
        st.markdown(f"**Transcription:** {transcript_behavior}")

        transcript_development = transcribe_audio(audio_development, name='development')
        st.markdown(f"**Transcription:** {transcript_development}")
        
        response = client.responses.create(
            prompt={
                "id": st.secrets["behavior_prompt_id"],
                # "version": "3",
                "variables": {
                    "first_name": data['{{Patient First Name}}'],
                    "pronouns": preferred,
                    "diagnosis": "having autism",
                    "transcription": transcript_behavior
                }
            }
        )
        st.session_state.behavior_observation = response.output_text
        # behavior_observation = response.output_text

        # calculate tokens
        st.write("Input tokens:", response.usage.input_tokens)
        st.write("Output tokens:", response.usage.output_tokens)

        response = client.responses.create(
            prompt={
                "id": st.secrets["development_prompt_id"],
                # "version": "5",
                "variables": {
                    "first_name": data['{{Patient First Name}}'],
                    "pronouns": preferred,
                    "diagnosis": "having autism",
                    "transcription": transcript_development
                }
            }
        )
        st.session_state.development_history = response.output_text
        # development_history = response.output_text
        
        # calculate tokens
        st.write("Input tokens:", response.usage.input_tokens)
        st.write("Output tokens:", response.usage.output_tokens)

with st.form('EditResponse'):
    st.header("Edit OpenAI Response")

    # st.markdown("**Behavioral Observation:**")
    data['behavior_observation'] = st.text_area(
        "Behavioral Observation: Edit the response before submitting the form", 
        # behavior_observation,
        st.session_state.behavior_observation,
        height=400,
    )

    # st.markdown("**Developmental History:**")
    data['development_history'] = st.text_area(
        "Developmental History: Edit the response before submitting the form", 
        # development_history,
        st.session_state.development_history,
        height=400,
    )
    
    data['{{Residence City/State}}'] = st.text_input("Residence City/State")
    # st.selectbox(
    #     "Residence City/State", states, index=None,
    # )

    data['{{Narrative}}'] = st.text_area('Narrative to finish \"Patient lives with...\"')
    
    submit = st.form_submit_button('Submit')

def format_date_with_ordinal(date_obj):
    day = date_obj.day
    suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return date_obj.strftime(f"%B {day}{suffix}, %Y")

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def add_behavior_presentation(paragraph, transcript):
    # separate transcript
    small_para = transcript.split('\n')

    st.write(small_para)

    paragraph.insert_paragraph_before().add_run(small_para[0], style='CustomStyle')
    paragraph.insert_paragraph_before()

    for sub_para in small_para[1:]:
        sub_para = sub_para.split(":")
        p = paragraph.insert_paragraph_before()
        p.add_run(sub_para[0], style='CustomStyle').italic = True
        p.add_run(f":{sub_para[1]}\n", style='CustomStyle')
        
    delete_paragraph(paragraph)

def add_developmental_history(paragraph, transcript):
    # separate transcript
    small_para = transcript.split('\n')
    st.write(small_para)

    for sub_para in small_para:
        sub_para = sub_para.split(":")
        p = paragraph.insert_paragraph_before()
        p.add_run(sub_para[0], style='CustomStyle').italic = True
        p.add_run(f":{sub_para[1]}\n", style='CustomStyle')
        
    delete_paragraph(paragraph)


if submit:
    # st.session_state.final_text = data['Transcription']
    # Display data 
    yaml_string = yaml.dump(data, sort_keys=False)
    yaml_data = st.code(yaml_string, language=None)

    #### Edit document 
    doc = Document('templates/template_mod_12.docx')
    if doc:
        # Get file name
        today_date = format_date_with_ordinal(datetime.date.today())
        filename = f"{data['{{Patient First Name}}']} {data['{{Patient Last Name}}']} {today_date}.docx"
        
        ### create document 
        norm_style = doc.styles['Normal']
        norm_style.paragraph_format.line_spacing = 1

        custom_style = doc.styles.add_style('CustomStyle', WD_STYLE_TYPE.CHARACTER)
        custom_style.font.size = Pt(12)
        custom_style.font.name = 'Georgia'
        # custom_style.paragraph_format.line_spacing = 1 

        list_style = doc.styles['Bullet New']
        list_style.paragraph_format.line_spacing = 1

        # Add scores 
        for i, paragraph in enumerate(doc.paragraphs):
            if "[[Behavioral Presentation]]" in paragraph.text:
                add_behavior_presentation(paragraph, st.session_state.behavior_observation)
            if "[[Developmental History]]" in paragraph.text:
                add_behavior_presentation(paragraph, st.session_state.behavior_observation)

        # Save content to file
        doc.save(filename)

        # Download 
        bio = io.BytesIO()
        document = Document(filename)
        document.save(bio)

        st.download_button(
            label="Click here to download",
            key="report_download",
            data=bio.getvalue(),
            file_name=filename,
            mime="docx"
        )
