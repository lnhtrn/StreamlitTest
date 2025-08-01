import streamlit as st
from docx import Document
import docx
import yaml
import io
import docxedit
import datetime
from docx.enum.style import WD_STYLE_TYPE
from streamlit_gsheets import GSheetsConnection
from docxtpl import DocxTemplate
from docx.shared import Inches, Pt
from openai import OpenAI
from modules.recommendations import *


def add_hyperlink(paragraph, url, size=24):
    """
    A function that places a hyperlink within a paragraph object with custom font and size.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :param color: Hex color string (e.g., '0000FF')
    :param underline: Bool indicating whether the link is underlined
    :return: The hyperlink object
    """

    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set font to Georgia
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Georgia')
    rFonts.set(qn('w:hAnsi'), 'Georgia')
    rPr.append(rFonts)

    # Set font size to 11.5pt (23 half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), f'{size}')
    rPr.append(sz)

    c = OxmlElement('w:color')
    c.set(qn('w:val'), '1155cc')
    rPr.append(c)

    # Set underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = url
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def add_bullet(paragraph, text):
    paragraph.style = 'Bullet New'
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.add_run(text, style='CustomStyle')
    
def add_bullet_and_link(paragraph, text, url):
    paragraph.style = 'Bullet New'
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.add_run(text, style='CustomStyle')
    add_hyperlink(paragraph, url)


def add_bold(paragraph, text_bold):
    r = paragraph.add_run(f"{text_bold} ", style='CustomStyle')
    r.bold = True
    r.italic = True
    
def add_normal(paragraph, text):
    paragraph.add_run(text, style='CustomStyle')


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


# set up recommendation system
check_rec = {}
with open("misc_data/rec_per_module.yaml", "r") as file:
    rec_all = yaml.safe_load(file)
    recommendation_options = {}
    for mod in rec_all:
        recommendation_options.update(rec_all[mod])

### Create new doc
doc = Document('templates/template_mod_12.docx')

### get styles
norm_style = doc.styles['Normal']
norm_style.paragraph_format.line_spacing = 1

custom_style = doc.styles.add_style('CustomStyle', WD_STYLE_TYPE.CHARACTER)
custom_style.font.size = Pt(12)
custom_style.font.name = 'Georgia'

custom_style_2 = doc.styles.add_style('CustomStyle2', WD_STYLE_TYPE.CHARACTER)
custom_style_2.font.size = Pt(11.5)
custom_style_2.font.name = 'Georgia'

list_style = doc.styles['Bullet New']
list_style.paragraph_format.line_spacing = 1

### add paragraphs

for paragraph in doc.paragraphs:
    if "[[Recommendations]]" in paragraph.text:
        for rec in recommendation_options:
            func = globals().get(f"add_{rec}")
            if callable(func):
                func(paragraph)
        delete_paragraph(paragraph)

doc.save('recommendation_test.docx')