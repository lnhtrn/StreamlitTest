import docx
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement, qn


###############################################################
# Universal
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def add_hyperlink(paragraph, url, size=24):
    """
    A function that places a hyperlink within a paragraph object with custom font and size.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :param color: Hex color string (e.g., '0000FF')
    :param underline: Bool indicating whether the link is underlined
    :return: The hyperlink object
    """

    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set font to Georgia
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Georgia')
    rFonts.set(qn('w:hAnsi'), 'Georgia')
    rPr.append(rFonts)

    # Set font size to 11.5pt (23 half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), f'{size}')
    rPr.append(sz)

    c = OxmlElement('w:color')
    c.set(qn('w:val'), '1155cc')
    rPr.append(c)

    # Set underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = url
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

###############################################################
# Recommendations Mod 1&2

def add_developmental_pediatrics(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Developmental Pediatrics Appointment. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('I believe that {{Patient First Name}} would benefit from being seen by a developmental medical provider as part of comprehensive care related to the diagnosis described here. An appointment can be made by calling one of the following local specialty clinics or at URMC and Rochester Regional Health Center:\n', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('University of Rochester Medical Center, Levine Autism Clinic at 585-275-2986,', style='CustomStyle2')
    p = paragraph.insert_paragraph_before(style='Normal')
    p.paragraph_format.left_indent = Inches(0.5)
    add_hyperlink(p, 'https:/www.urmc.rochester.edu/childrens-hospital/developmental-disabilities/services/levine.aspx', size=23)

    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Rochester Regional Health Center, Developmental Behavioral Pediatrics Program at 585-922-4698, ', style='CustomStyle2')
    add_hyperlink(p, 'https://www.rochesterregional.org/services/pediatrics/developmental-behavioral-pediatrics-program', size=23)
    paragraph.insert_paragraph_before()

def add_feeding_treatment(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Feeding Treatment & Support. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('{{Patient First Name}} presents with a range of concerns related to mealtime behavior and food variety, so I recommend that {{Preferred Pronouns 2}} parents seek out support from one of the following local agencies. I am happy to discuss this in detail.\n', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('University of Rochester Medical Center - ', style='CustomStyle2')
    p = paragraph.insert_paragraph_before(style='Normal')
    p.paragraph_format.left_indent = Inches(0.5)
    add_hyperlink(p, 'https://www.urmc.rochester.edu/childrens-hospital/developmental-disabilities/services/feeding-disorders.aspx')

    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Step-by-Step - ', style='CustomStyle')
    add_hyperlink(p, 'https://www.sbstherapycenter.com/feeding-therapy')
    
    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Mealtime Rediscovered - ', style='CustomStyle')
    add_hyperlink(p, 'https://mealtimerediscovered.com/')
    paragraph.insert_paragraph_before()

def add_levine_clinic(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Levine Autism Clinic. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('I recommend that {{Patient First Name}}’s {{Caregiver type}} refer to the Levine Autism Clinic Facebook page for information about services, supports, events, and information that may be of help: ', style='CustomStyle')
    p = paragraph.insert_paragraph_before(style='Normal')
    add_hyperlink(p, 'https://www.facebook.com/DBPeds.GCH/')
    paragraph.insert_paragraph_before()

def add_parent_parent(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Parent to Parent. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('(', style='CustomStyle')
    add_hyperlink(p, 'http://parenttoparentnys.org/offices/Finger-Lakes/')
    p.add_run(') This group could help to connect {{Patient First Name}}’s family with another family in their area who knows more about local resources and supports related to {{Patient First Name}}’s age-level and interests.', style='CustomStyle')
    paragraph.insert_paragraph_before()

def add_100_days(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Autism Speaks 100 Days 100 Kit. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('I would recommend that {{Patient First Name}}’s {{Caregiver type}} refer to this kit to help structure their next steps in determining {{Patient First Name}}’s care. The kit contains information and advice collected from trusted and respected experts. ', style='CustomStyle')
    p = paragraph.insert_paragraph_before(style='Normal')
    add_hyperlink(p, 'http://www.autismspeaks.org/community/family_services/100_day_kit.php')
    paragraph.insert_paragraph_before()

def add_caregiver_support(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Caregiver Support.  ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('I encourage {{Patient First Name}}’s {{Caregiver type}} to review these resources:\n', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('AutismUp - ', style='CustomStyle')
    add_hyperlink(p, 'https://autismup.org/support/family-navigator')

    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Autism Council of Rochester - ', style='CustomStyle')
    add_hyperlink(p, 'https://www.theautismcouncil.org/')

    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Camp Puzzle Peace - ', style='CustomStyle')
    add_hyperlink(p, 'www.familyautismcenter.com/')

    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Rochester Regional Center for Autism Spectrum Disorders - \n', style='CustomStyle')
    p = paragraph.insert_paragraph_before(style='Normal')
    p.paragraph_format.left_indent = Inches(0.5)
    add_hyperlink(p, 'https://www.urmc.rochester.edu/strong-center-developmental-disabilities/programs/rochester-regional-ctr-autism-spectrum-disorder.aspx')

    paragraph.insert_paragraph_before()

def add_edu_placement(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Educational Placement. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('The matter of which setting {{Patient First Name}} is educated in feels of paramount concern given {{Preferred Pronouns 2}} current skills and areas of need. I encourage {{Preferred Pronouns 2}} {{Caregiver type}} and school team to engage in ongoing conversations about placement options available for next year. I recommend that discussions about educational placement and programming be held within the CPSE meeting process.', style='CustomStyle')
    p = paragraph.insert_paragraph_before()

def add_effective_treatments(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Components of Effective Treatment. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('A body of research has accumulated about effective treatment for children with autism. A list of components of this presented below. How these are implemented is best determined by those who work with {{Patient First Name}}. \n', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Comprehensive curriculum focusing on teaching a wide range of skills, including attention to the environment, imitation, comprehension and production of language, functional communication, toy play, and peer interaction.', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Supportive teaching environments structured to maximize attention to tasks.', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Emphasis on providing children with predictability and routine.', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Functional behavior analytic approach to assessing and treating behaviors.', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Systematic intervention for facilitating transitions from home to school setting.', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Consultation with a professional with expertise in autism-related interventions.', style='CustomStyle')

    paragraph.insert_paragraph_before()

def add_elopement_plan(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Elopement Plan. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('Given {{Patient First Name}}’s predisposition to wander and bolt if not closely monitored, I think that it is medically necessary for {{Preferred Pronouns 2}} team to have in place a series of preventative and responsive procedures related to {{Preferred Pronouns 2}} elopement. This could be done in consultation with the school team (teacher, social worker) and a behavior specialist.\nResources to consider include:\n', style='CustomStyle')
    
    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Big Red Safety Toolkit - ', style='CustomStyle')
    p = paragraph.insert_paragraph_before(style='Normal')
    p.paragraph_format.left_indent = Inches(0.5)
    add_hyperlink(p, 'https://nationalautismassociation.org/docs/BigRedSafetyToolkit.pdf')

    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Angel Sense - ', style='CustomStyle')
    add_hyperlink(p, 'https://www.angelsense.com/gps-tracker-lifesaving-features/')

    paragraph.insert_paragraph_before(style='Normal')

def add_develop_disability_office(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Developmental Disabilities Regional Office (DDRO). ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('I discussed DDRO case management and Medicaid Waiver services with {{Patient First Name}}’s {{Caregiver type}}. To qualify for services, a person must have a diagnosis of a developmental disability along with documentation of cognitive and/or adaptive deficits. Based on {{Preferred Pronouns 2}} presentation and chart review, I believe that {{Patient First Name}} ought to quality for OPWDD waiver services due to {{Preferred Pronouns 2}} adaptive and cognitive delays. More information on Front Door Sessions can be found online at: ', style='CustomStyle')
    add_hyperlink(p, 'https://opwdd.ny.gov/get-started/information-sessions')
    paragraph.insert_paragraph_before()
    
    p = paragraph.insert_paragraph_before()
    p.add_run('Information can be obtained through the Office of Persons with Developmental Disabilities (OPWDD), ', style='CustomStyle')
    p.add_run('Front Door Office Finger Lakes', style='CustomStyle').bold = True
    p.add_run(' at 855-679-3335', style='CustomStyle')
    paragraph.insert_paragraph_before()

def add_evidence_therapy(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Evidence-Based Therapies. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('I would encourage {{Patient First Name}}’s family to consider seeking services that are informed by the principles of applied behavior analysis (ABA). In particular, I would recommend that {{Patient First Name}} receive intensive intervention under the supervision of a licensed professional or board-certified behavioral analyst.\n\nResources to consider include:\n', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Autism Learning Partners - ', style='CustomStyle')
    add_hyperlink(p, 'https://www.autismlearningpartners.com/')

    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Living Soul - ', style='CustomStyle')
    add_hyperlink(p, 'https://livingsoulllc.com/')

    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Proud Moments - ', style='CustomStyle')
    add_hyperlink(p, 'https://discover.proudmomentsaba.com/rochester.html')

    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('TruNorth Autism Services - \n', style='CustomStyle')
    add_hyperlink(p, 'https://www.trunorthautism.com/')

    paragraph.insert_paragraph_before()




# Recommedation Mod 3 No 

def add_accomodations(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Accommodations. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('I would recommend that a designated location be determined for assessments, especially those with a timed component. There is also value providing {{Patient First Name}} with additional time and supports to remember to ask for accommodations.\n', style='CustomStyle')

def add_check_in_anxiety(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Check-in Support for Anxiety. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('{{Patient First Name}}’s parents reported that {{Preferred Pronouns 1}} has an anxiety diagnosis, and we would recommend adding counseling supports to help determine if anxiety is contributing to difficulties with auditory filtering or attention on assessments.\n', style='CustomStyle')

def add_hearing_assistive(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Hearing Assistive Technology. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run("Given {{Patient First Name}}'s presentation, results of assessments, and behavioral observations across settings, we recommend trialing the use of a hearing assistive technology to address concerns related to auditory filtering and attention.\n", style='CustomStyle')

def add_additional_reading_suggestions(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Additional Suggestions. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run("Given {{Patient First Name}}'s reading difficulties, I recommend the following:\n", style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Enlarged font with strategic spacing between lines of text for assignments\n', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Hearing assistive technology to rule-out potential auditory or attention concerns\n', style='CustomStyle')

    p = paragraph.insert_paragraph_before()
    p.add_run("In so far as {{Patient First Name}}'s difficulties could be attributed to, in part, to auditory processing or filtering concerns, I think there’s value in trialing the use of a hearing assistive tech.\n", style='CustomStyle').italic = True
    
def add_home_community(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Home/Community Based Supports. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('Based on the report by {{Preferred Pronouns 2}} mother and our observations today, we recommend considering therapeutic supports through organizations that provide specialized care. To that end, we recommend the following organizations: \n', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('URMC Pediatric Behavioral Health and Wellness:\n', style='CustomStyle')
    p = paragraph.insert_paragraph_before(style='Normal')
    p.paragraph_format.left_indent = Inches(0.5)
    add_hyperlink(p, 'https://www.urmc.rochester.edu/childrens-hospital/behavioral-health-wellness/outpatient')
    p = paragraph.insert_paragraph_before()

def add_coordinated_care(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Coordinated Care. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('We encourage collaboration between home, school, and medical providers to ensure a unified approach to supporting {{Patient First Name}}’s needs across settings.\n', style='CustomStyle')

def add_therapeutic_supports(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Therapeutic Supports. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('Based on the report by {{Preferred Pronouns 2}} mother and our observations today, we recommend considering therapeutic supports through organizations that provide specialized care. To that end, we recommend the following organizations: \n', style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('URMC Pediatric Behavioral Health and Wellness:\n', style='CustomStyle')
    p = paragraph.insert_paragraph_before(style='Normal')
    p.paragraph_format.left_indent = Inches(0.5)
    add_hyperlink(p, 'https://www.urmc.rochester.edu/childrens-hospital/behavioral-health-wellness/outpatient')

    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run('Genesee Valley Psychology: ', style='CustomStyle')
    add_hyperlink(p, 'https://gviproc.org/about')
    paragraph.insert_paragraph_before()

def add_coordinated_healthcare(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Therapeutic Supports. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run("Based on {{Patient First Name}}'s history of mental health and inter-related social concerns, we think there is value in pursuing coordinated healthcare such as:\n", style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    add_hyperlink(p, 'https://www.rochesterregional.org/services/adult-mental-health/pros')

    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before(style='Bullet New')
    add_hyperlink(p, 'https://www.hhuny.org/')
    paragraph.insert_paragraph_before()

def add_executive_functioning(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Executive Functioning Accommodations. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run("{{Patient First Name}} will need frequent check-ins to review {{Preferred Pronouns 2}} daily and weekly activities to make sure {{Preferred Pronouns 1}} remains on target with longer-term assignments. Organization of materials is a critical component of this process, and, as such, we recommend daily adult support to initiate and complete work for {{Patient First Name}}.\n\nOne approach that might be worth pursuing at school would be:\n", style='CustomStyle')

    p = paragraph.insert_paragraph_before(style='Bullet New')
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run("Unstuck & On Target - ", style='CustomStyle')
    p = paragraph.insert_paragraph_before()
    p.paragraph_format.left_indent = Inches(0.5)
    add_hyperlink(p, 'https://www.rochesterregional.org/services/adult-mental-health/pros')
    p = paragraph.insert_paragraph_before()

def add_support(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Support. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('We think that {{Patient First Name}} would benefit from the school’s behavioral specialist to accommodate {{Preferred Pronouns 2}} day-to-day management needs at school. We encourage opportunities to engage with {{Preferred Pronouns 2}} peers as well as executive functioning suggestions outlined above.\n', style='CustomStyle')

def add_adult_career(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Adult Career and Continuing Education Services-Vocational Rehabilitation (ACCES-VR)', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run('{{Patient First Name}} may benefit from vocational support to help maintain employment and supported {{Preferred Pronouns 2}} independent living. Information can be found on the NYSD Department’s website: ', style='CustomStyle')
    add_hyperlink(p, 'https://www.acces.nysed.gov/vr')
    paragraph.insert_paragraph_before()

def add_occupational_therapy(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Occupational Therapy. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run("Based on {{Patient First Name}}'s parents’ report, there is a clear pattern of concerns related to auditory aversions that set to the occasion for potential elopement. As such, I would recommend reaching out to URMC Pediatric Occupational Therapy:\n", style='CustomStyle')
    
    p = paragraph.insert_paragraph_before(style='Bullet New')
    add_hyperlink(p, 'https://www.urmc.rochester.edu/locations/pediatric-neurology')
    p = paragraph.insert_paragraph_before()

def add_pediatric_neurology(paragraph):
    p = paragraph.insert_paragraph_before()
    r = p.add_run('Pediatric Neurology Consultation. ', style='CustomStyle')
    r.bold = True
    r.italic = True
    p.add_run("I observed {{Patient First Name}} to move {{Preferred Pronouns 2}} face and eyes in a repetitive manner that appeared involuntary; I noted repetitive vocalization of sounds and potential throat-clearing. Based on this, I recommend that {{Patient First Name}}'s parents consult with URMC Pediatric Neurology to determine if these behaviors reflect specific tics or Tourette's syndrome. More information can be found here can be found by calling 585-275-2808 or going to \n", style='CustomStyle')
    add_hyperlink(p, 'https://www.urmc.rochester.edu/locations/pediatric-neurology')
    p = paragraph.insert_paragraph_before()
