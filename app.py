import streamlit as st
from docx import Document
import yaml
import io
import docxedit
import datetime
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from streamlit_gsheets import GSheetsConnection
from docxtpl import DocxTemplate

primary_concerns = []

# Create a connection object for google sheets
# def load_data(store_data):
#     conn = st.connection("gsheets", type=GSheetsConnection)
#     store_data = conn.read(worksheet="Sheet1")['Results'].tolist()
# st.button("Reload Data", on_click=load_data(primary_concerns))

col1,col2 = st.columns(2)
col1.title('Report Builder')

with open('misc_data/states.txt', 'r') as file:
    states = [x.strip() for x in file]

def format_date_with_ordinal(date_obj):
    day = date_obj.day
    suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return date_obj.strftime(f"%B {day}{suffix}, %Y")

# Set up dictionary to store data 
data = {}
optional = {}
teacher_score = {}
bullet = {}
lines = {}
comma = {}

####################################################
st.header("Appointment Summary")
audio_summary = st.audio_input("Summarize the appointment details")
st.markdown("***Check before proceeding with form:*** Scores to report:")
teacher_eval = st.checkbox("Teacher's SSR Scores")
wppsi_score = st.checkbox("Wechsler Preschool & Primary Scales of Intelligence – Fourth Ed. (WPPSI) Score")
dppr_score = st.checkbox("Developmental Profile – Fourth Edition – Parent Report (DPPR)")
pls_score = st.checkbox("Preschool Language Scale – Fifth Edition (PLS)")
pdms_score = st.checkbox("Peabody Developmental Motor Scales – Second Edition")
peshv_score = st.checkbox("Preschool Evaluation Scale Home Version – Second Edition")
reelt_score = st.checkbox("Receptive Expressive Emergent Language Test – Fourth Edition")
abas_score = st.checkbox("Adaptive Behavior Assessment System – Third Edition")

with st.form('BasicInfo'):
    ####################################################
    st.header("Patient's data")
    # Dict to store data
    data['{{Patient First Name}}'] = st.text_input('Patient First Name')
    data['{{Patient Last Name}}'] = st.text_input('Patient Last Name')
    preferred = st.selectbox(
        "Patient's Preferred Pronoun",
        ("They/them", "He/him", "She/her"),
    )
    data["{{Patient Age}}"] = st.number_input("Patient's Age", 0, 100)
    data['{{Patient age unit}}'] = st.radio(
        "Year/month?",
        ("year", "month")
    )

    data['{{Caregiver type}}'] = st.selectbox(
        "Patient's Caregiver",
        ("mother", "father", "parent", "grandparent", "legal custodian", "foster parent"),
        placeholder="Select from the choices or enter a new one",
        index=None,
        accept_new_options=True,
    )

    bullet['Caregiver_Primary_Concerns'] = st.multiselect(
        "Caregiver\'s Primary Concerns",
        [
            "Speech delays impacting social opportunities.",
            "Clarifying diagnostic presentation.",
            "Determining service eligibility.",
            "Language delays and difficulties.",
            "Elopement and related safety concerns.",
            "Determining appropriate supports."
        ],
        placeholder="Select from the choices or enter a new one",
        accept_new_options=True
    )
    
    data['{{Residence City/State}}'] = st.text_input("Residence City/State")
    # st.selectbox(
    #     "Residence City/State", states, index=None,
    # )

    data['{{Narrative}}'] = st.text_area('Narrative to finish \"Patient lives with...\"')

    ##########################################################
    st.header("BRH Evaluation Details")

    data['{{Evaluation Date}}'] = format_date_with_ordinal(st.date_input("Evaluation Date"))

    data['{{Module used}}'] = st.radio("Module used", ["Module 1", "Module 2"])
    if data['{{Module used}}'] == "Module 1":
        data['{{Module Description}}'] = "Module 1 is designed for children with single words"
    else:
        data['{{Module Description}}'] = "Module 2 is designed for children with phrase speech"

    data['{{Location of the evaluation}}'] = st.radio(
        "Location of the evaluation",
        ['home', 'school', 'the office'],
        index=None,
    )

    data['{{Results Shared Date}}'] = format_date_with_ordinal(st.date_input("Results Shared Date"))
    
    data['{{Date Report Sent to Patient}}'] = format_date_with_ordinal(st.date_input("Date Report Sent to Patient"))

    lines["{{Result of the evaluation}}"] = st.multiselect(
        "Result of the evaluation",
        [
            "F84.0 - Autism Spectrum Disorder (per the above referenced evaluation)",
            "F88.0 - Global Developmental Delay (per behavioral presentation)",
            "F80.2 - Mixed Receptive-Expressive Language Disorder",
            "F90.2 - Attention Deficit Hyperactivity Disorder - Combined-Type",
            "F50.82 Avoidant/Restrictive Food Intake Disorder",
            "None"
        ],
        placeholder="Select from the choices or enter a new one",
        accept_new_options=True
    )

    data["{{Results (SCQ) - Lifetime Form}}"] = st.text_input(
        "Results (SCQ) - Lifetime Form"
    )

    data["{{SRS-2 Score Caregiver}}"] = st.text_input("Caregiver's SRS-2 Score")
    
    data["{{Social Communication and Interaction Score Caregiver}}"] = st.text_input("Social Communication and Interaction Score Caregiver")
    
    data["{{Restricted Interests and Repetitive Behavior Score Caregiver}}"] = st.text_input("Restricted Interests and Repetitive Behavior Score Caregiver")

    data["{{Caregiver's level of concern}}"] = st.radio(
        "Caregiver's level of concern",
        ['no', 'mild', 'moderate', 'severe']
    )

    data["{{Evaluator's level of concern}}"] = st.radio(
        "Evaluator's level of concern",
        ['no', 'mild', 'moderate', 'severe']
    )

    ##########################################################
    if teacher_eval:
        st.header("Teacher SSR Score")
        st.markdown("*Skip this section if teacher did not give SSR Score*")

        teacher_score['{{SRS-2 Score Teacher}}'] = st.text_input("Teacher's SRS-2 Score")

        teacher_score['{{Social Communication and Interaction Score Teacher}}'] = st.text_input("Social Communication and Interaction Score Teacher")

        teacher_score['{{Restricted Interests and Repetitive Behavior Score Teacher}}'] = st.text_input("Restricted Interests and Repetitive Behavior Score Teacher")

        teacher_score["{{Teacher level of concern}}"] = st.radio(
            "Teacher's level of concern",
            ['no', 'mild', 'moderate', 'severe']
        )

    ######################################################
    st.header("Medical/Developmental History")
    
    lines['{{Diagnosis History}}'] = st.multiselect(
        "Diagnosis History (Select or add your own)",
        ['History of language and social communication delays.'],
        accept_new_options=True
    )

    lines['{{Medications}}'] = st.multiselect(
        "Medications (Select or add your own)",
        ['None noted or reported.'],
        accept_new_options=True
    )

    ###############################################
    st.header("Educational Background")

    data['{{School District}}'] = st.selectbox(
        "School District",
        ['Rochester City'],
        index=None,
        placeholder="Select a school district or enter a new one",
        accept_new_options=True,
    )

    data['{{School Name}}'] = st.text_input("School Name")

    data['{{Grade}}'] = st.selectbox(
        "Grade",
        ['EPK (2023-24 school year)', 'UPK (2023-24 school year)', 'Kindergarten (2023-24 school year)'],
        index=None,
        placeholder="Select a grade or enter a new one",
        accept_new_options=True,
    )

    data['{{Teacher name, title}}'] = st.text_input("Teacher name, title")

    data['{{Education Setting}}'] = st.selectbox(
        "Education Setting",
        ["General Education", "Integrated Co-Taught", "12:1:1", "8:1:1", "6:1:1"],
        index=None,
        placeholder="Select a grade or enter a new one",
        accept_new_options=True,
    )

    comma['{{Services}}'] = st.multiselect(
        "Services",
        [
            "None",
            "Speech therapy",
            "Occupational therapy",
            "Physical therapy",
            "Extended school year services",
            "Testing accommodations"
        ],
        placeholder="Select from the choices or enter a new one",
        accept_new_options=True
    )

    ##########################################################
    if wppsi_score:
        st.header("Wechsler Preschool & Primary Scales of Intelligence – Fourth Ed. (WPPSI)")
        st.markdown("*Skip this section if there is no WPPSI Score*")
        optional["wppsi"] = {}

        optional["wppsi"]["Test Date"] = st.date_input("WPPSI Test Date").strftime("%m/%Y")
        optional["wppsi"]['WPPSI Full Scale IQ Score'] = st.text_input("WPPSI Full Scale IQ Score")

        optional["wppsi"]['WPPSI Verbal Comprehension Score'] = st.text_input("WPPSI Verbal Comprehension Score")

        optional["wppsi"]['WPPSI Visual Spatial Score'] = st.text_input("WPPSI Visual Spatial Score")
    
    if dppr_score:
        st.header("Developmental Profile – Fourth Edition – Parent Report (DPPR)")
        st.markdown("*Skip this section if there is no DPPR Score*")
        optional["dppr"] = {}

        optional["dppr"]["Test Date"] = st.date_input("DPPR Test Date").strftime("%m/%Y")
        optional["dppr"]['DPPR Cognitive Score'] = st.text_input("DPPR Cognitive Score")
        optional["dppr"]['DPPR Social-Emotional Score'] = st.text_input("DPPR Social-Emotional Score")
        optional["dppr"]['DPPR Adaptive Score'] = st.text_input("DPPR Adaptive Score")
        optional["dppr"]['DPPR Physical Score'] = st.text_input("DPPR Physical Score")
    
    if pls_score:
        st.header("Preschool Language Scale – Fifth Edition (PLS)")
        st.markdown("*Skip this section if there is no PLS Score*")
        optional["pls"] = {}
        optional["pls"]["Test Date"] = st.date_input("PLS Test Date").strftime("%m/%Y")
        optional["pls"]['PLS Total Language Score'] = st.text_input("PLS Total Language Score")
        optional["pls"]['PLS Auditory Comprehension Score'] = st.text_input("PLS Auditory Comprehension Score")
        optional["pls"]['PLS Expressive Communication Score'] = st.text_input("PLS Expressive Communication Score")

    if pdms_score:
        st.header("Peabody Developmental Motor Scales – Second Edition (PDMS)")
        st.markdown("*Skip this section if there is no PDMS Score*")
        optional["pdms"] = {}
        optional["pdms"]["Test Date"] = st.date_input("Test Date").strftime("%m/%Y")
        optional["pdms"]['PDMS Gross Motor Score'] = st.text_input("PDMS Gross Motor Score")
        optional["pdms"]['PDMS Fine Motor Score'] = st.text_input("PDMS Fine Motor Score")

    if peshv_score:
        st.header("Preschool Evaluation Scale Home Version – Second Edition (PESHV)")
        st.markdown("*Skip this section if there is no PESHV Score*")
        optional["peshv"] = {}
        optional["peshv"]["Test Date"] = st.date_input("PESHV Test Date").strftime("%m/%Y")
        optional["peshv"]['PESHV Cognitive Score'] = st.text_input("PESHV Cognitive Score")
        optional["peshv"]['PESHV Social Emotional Score'] = st.text_input("PESHV Social Emotional Score")
    
    if peshv_score:
        st.header("Receptive Expressive Emergent Language Test – Fourth Edition (REELT)")
        st.markdown("*Skip this section if there is no REELT Score*")
        optional[""] = {}
        optional["peshv"]["Test Date"] = st.date_input("PESHV Test Date").strftime("%m/%Y")
        optional["peshv"]['Total Language'] = st.text_input("Total Language")
        optional["peshv"]['PESHV Social Emotional Score'] = st.text_input("PESHV Social Emotional Score")

    if reelt_score:
        st.header("Receptive Expressive Emergent Language Test – Fourth Edition (REELT)")
        st.markdown("*Skip this section if there is no REELT Score*")
        optional["reelt"] = {}
        optional["reelt"]["Test Date"] = st.date_input("REELT Test Date").strftime("%m/%Y")
        optional["reelt"]['REELT Total Language Score'] = st.text_input("REELT Total Language Score")
        optional["reelt"]['REELT Auditory Comprehension Score'] = st.text_input("REELT Auditory Comprehension Score")
        optional["reelt"]['REELT Expressive Communication Score'] = st.text_input("REELT Expressive Communication Score")

    if abas_score:
        st.header("Adaptive Behavior Assessment System – Third Edition (ABAS)")
        st.markdown("*Skip this section if there is no ABAS Score*")
        optional["abas"] = {}
        optional["abas"]["Test Date"] = st.date_input("ABAS Test Date").strftime("%m/%Y")
        optional["abas"]['ABAS General Adaptive Composite'] = st.text_input("ABAS General Adaptive Composite")
        optional["abas"]['ABAS Conceptual'] = st.text_input("ABAS Conceptual")
        optional["abas"]['ABAS Social'] = st.text_input("ABAS Social")
        optional["abas"]['ABAS Practical'] = st.text_input("ABAS Practical")

    ############################################
    st.header("DSM Criteria")
    
    bullet['Deficits_in_social_emotional_reciprocity'] = st.multiselect(    "Deficits in social emotional reciprocity",
        [
            "None",
            "Awkward social initiation and response",
            "Difficulties with chit-chat",
            "Difficulty interpreting figurative language",
            "Limited social approach or greetings",
        ],
        placeholder="Select from the choices or enter a new one",
        accept_new_options=True
    )

    bullet['Deficits_in_nonverbal_communicative_behaviors_used_for_social_interaction'] = st.multiselect(
        "Deficits in nonverbal communicative behaviors used for social interaction",
        [
            "None",
            "Limited well-directed eye contact",
            "Difficulty reading facial expressions",
            "Absence of joint attention",
            "Lack of well-integrated gestures",
            "Limited range of facial expression",
        ],
        placeholder="Select from the choices or enter a new one",
        accept_new_options=True
    )

    bullet['Deficits_in_developing_maintaining_and_understanding_relationships'] = st.multiselect(
        "Deficits in developing, maintaining, and understanding relationships",
        [
            "None",
            "Limited engagement with same age peers",
            "Difficulties adjusting behavior to social context",
            "Difficulties forming friendships",
        ],
        placeholder="Select from the choices or enter a new one",
        accept_new_options=True
    )

    bullet['Stereotyped_or_repetitive_motor_movements_use_of_objects_or_speech'] = st.multiselect(
        "Stereotyped or repetitive motor movements, use of objects, or speech",
        [
            "None",
            "Repetitive whole-body movements",
            "Repetitive hand movements",
            "Echolalia of sounds",
            "Echolalia of words",
            "Stereotyped speech",
        ],
        placeholder="Select from the choices or enter a new one",
        accept_new_options=True
    )

    bullet['Insistence_on_sameness_inflexible_adherence_to_routines_or_ritualized_behavior'] = st.multiselect(
        "Insistence on sameness, inflexible adherence to routines or ritualized behavior",
        [
            "None",
            "Difficulties with changes in routine across developmental course",
            "Notable difficulties with transitions",
            "Insistence on following very specific routines",
        ],
        placeholder="Select from the choices or enter a new one",
        accept_new_options=True
    )

    bullet['Highly_restricted_fixated_interests_that_are_abnormal_in_intensity_or_focus'] = st.multiselect(
        "Highly restricted, fixated interests that are abnormal in intensity or focus",
        [
            "None",
            "Persistent pattern of perseverative interests",
            "Notable interest in topics others may find odd",
            "Very restricted pattern of eating and sleep time behavior",
        ],
        placeholder="Select from the choices or enter a new one",
        accept_new_options=True
    )

    bullet['Hyper_or_hypo_reactivity_to_sensory_aspects_of_the_environment'] = st.multiselect(
        "Hyper- or hypo-reactivity to sensory aspects of the environment:",
        [
            "None",
            "Auditory sensitivities",
            "Tactile defensiveness",
            "Proprioceptive-seeking behavior",
        ],
        placeholder="Select from the choices or enter a new one",
        accept_new_options=True
    )

    comma['{{Symptoms present in the early developmental period}}'] = st.multiselect(
        "Symptoms present in the early developmental period",
        [
            "Confirmed by record review",
            "None",
        ],
        placeholder="Select from the choices or enter a new one",
        accept_new_options=True
    )

    comma['{{Symptoms cause clinically significant impairment}}'] = st.multiselect(
        "Symptoms cause clinically significant impairment",
        [
            "Confirmed by record review",
            "None",
        ],
        placeholder="Select from the choices or enter a new one",
        accept_new_options=True
    )


    # data['{{}}'] = st.text_input("")
    # data['{{}}'] = st.text_input("")
    # data['{{}}'] = st.text_input("")
    # data['{{}}'] = st.text_input("")
    # data['{{}}'] = st.text_input("")
    # data['{{}}'] = st.text_input("")

    submit = st.form_submit_button('Submit')

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def add_srs_no_teacher(paragraph):
    r = paragraph.insert_paragraph_before().add_run('Social Responsiveness Scale – Second Edition (SRS-2) – Parent', style='CustomStyle')
    r.italic = True
    r.font.underline = True
    paragraph.insert_paragraph_before().add_run('The SRS-2 is an objective measure that identifies social impairments associated with autism spectrum disorder and quantifies ASD-related severity throughout the lifespan. \nThe following interpretative guidelines are offered here for the benefit of the reader: Less than 59 indicates within normal limits, between 60 and 65 as mild concern, between 65 and 75 as moderate concern, and greater than 76 as severe concern. ', style='CustomStyle')
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run('\tSRS-2 Total Score: {{SRS-2 Score Caregiver}} ({{Caregiver type}})', style='CustomStyle').bold = True
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run('\tSocial Communication and Interaction: {{Social Communication and Interaction Score Caregiver}} ({{Caregiver type}})', style='CustomStyle')
    paragraph.insert_paragraph_before().add_run('\tRestricted Interests and Repetitive Behavior: {{Restricted Interests and Repetitive Behavior Score Caregiver}} ({{Caregiver type}})', style='CustomStyle')
    paragraph.insert_paragraph_before()
    observe = paragraph.insert_paragraph_before()
    observe.add_run("Based on the report provided by {{Preferred Pronouns 2}} {{Caregiver type}}, ", style='CustomStyle')
    observe.add_run("{{Patient First Name}}’s social communication and related behaviors indicated {{Caregiver's level of concern}} concerns. ", style='CustomStyle').italic = True
    observe.add_run("My observation aligned with a {{Evaluator's level of concern}} level of concern.", style='CustomStyle').bold = True
    delete_paragraph(paragraph)

def add_srs_yes_teacher(paragraph, score_data):
    r = paragraph.insert_paragraph_before().add_run('Social Responsiveness Scale – Second Edition (SRS-2) – Parent', style='CustomStyle')
    r.italic = True
    r.font.underline = True
    paragraph.insert_paragraph_before().add_run('The SRS-2 is an objective measure that identifies social impairments associated with autism spectrum disorder and quantifies ASD-related severity throughout the lifespan. \nThe following interpretative guidelines are offered here for the benefit of the reader: Less than 59 indicates within normal limits, between 60 and 65 as mild concern, between 65 and 75 as moderate concern, and greater than 76 as severe concern. ', style='CustomStyle')
    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before()
    p.add_run('\tSRS-2 Total Score: {{SRS-2 Score Caregiver}} ({{Caregiver type}}), ', style='CustomStyle').bold = True
    p.add_run(f"{score_data['{{SRS-2 Score Teacher}}']} (teacher)", style='CustomStyle').bold = True
    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before()
    p.add_run('\tSocial Communication and Interaction: {{Social Communication and Interaction Score Caregiver}} ({{Caregiver type}}), ', style='CustomStyle')
    p.add_run(f"{score_data['{{Social Communication and Interaction Score Teacher}}']} (teacher)", style='CustomStyle')
    p = paragraph.insert_paragraph_before()
    p.add_run('\tRestricted Interests and Repetitive Behavior: {{Restricted Interests and Repetitive Behavior Score Caregiver}} ({{Caregiver type}}), ', style='CustomStyle')
    p.add_run(f'{score_data["{{Restricted Interests and Repetitive Behavior Score Teacher}}"]} (teacher)', style='CustomStyle')
    paragraph.insert_paragraph_before()
    observe = paragraph.insert_paragraph_before()
    observe.add_run("Based on the report provided by {{Preferred Pronouns 2}} {{Caregiver type}}, ", style='CustomStyle')
    observe.add_run("{{Patient First Name}}’s social communication and related behaviors indicated {{Caregiver's level of concern}} concerns. ", style='CustomStyle').italic = True
    observe.add_run("{{Patient First Name}}’s teacher reported a ", style='CustomStyle')
    observe.add_run(f"{score_data['{{Teacher level of concern}}']} level of concern, and ", style='CustomStyle')
    observe.add_run("my observation aligned with a {{Evaluator's level of concern}} level of concern.", style='CustomStyle').bold = True
    delete_paragraph(paragraph)

def add_wppsi(paragraph, score_data):
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run(f'\t({score_data["Test Date"]}) – Wechsler Preschool & Primary Scales of Intelligence – Fourth Ed.', style='CustomStyle').italic = True
    paragraph.insert_paragraph_before().add_run(f'\tFull Scale IQ: {score_data["WPPSI Full Scale IQ Score"]}', style='CustomStyle').bold = True
    paragraph.insert_paragraph_before().add_run(f'\tVerbal Comprehension: {score_data["WPPSI Verbal Comprehension Score"]}\t\t\tVisual Spatial: {score_data["WPPSI Visual Spatial Score"]}', style='CustomStyle')
    
def add_dppr(paragraph, score_data):
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run(f'\t({score_data["Test Date"]}) – Developmental Profile – Fourth Edition – Parent Report', style='CustomStyle').italic = True
    paragraph.insert_paragraph_before().add_run(f'\tCognitive: {score_data["DPPR Cognitive Score"]}\t\t\t\t\tSocial-Emotional: {score_data["DPPR Social-Emotional Score"]}', style='CustomStyle')
    paragraph.insert_paragraph_before().add_run(f'\tAdaptive: {score_data["DPPR Adaptive Score"]}\t\t\t\t\tPhysical: {score_data["DPPR Physical Score"]}', style='CustomStyle')

def add_pls(paragraph, score_data):
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run(f'\t({score_data["Test Date"]}) – Preschool Language Scale – Fifth Edition', style='CustomStyle').italic = True
    paragraph.insert_paragraph_before().add_run(f'\tTotal Language Score: {score_data["PLS Total Language Score"]}', style='CustomStyle').bold = True
    paragraph.insert_paragraph_before().add_run(f'\tAuditory Comprehension: {score_data["PLS Auditory Comprehension Score"]} \t\tExpressive Communication: {score_data["PLS Expressive Communication Score"]}', style='CustomStyle')

def add_pdms(paragraph, score_data):
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run(f'\t({score_data["Test Date"]}) – Peabody Developmental Motor Scales – Second Edition', style='CustomStyle').italic = True
    paragraph.insert_paragraph_before().add_run(f'\tGross Motor: {score_data["PDMS Gross Motor Score"]}\t\t\t\tFine Motor: {score_data["PDMS Fine Motor Score"]}', style='CustomStyle')
    
def add_peshv(paragraph, score_data):
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run(f'\t({score_data["Test Date"]}) – Preschool Evaluation Scale Home Version – Second Edition', style='CustomStyle').italic = True
    paragraph.insert_paragraph_before().add_run(f'\tCognitive: {score_data["PESHV Cognitive Score"]} \t\t\t\t\tSocial Emotional: {score_data["PESHV Social Emotional Score"]}', style='CustomStyle')

def add_reelt(paragraph, score_data):
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run(f'\t({score_data["Test Date"]}) – Receptive Expressive Emergent Language Test – Fourth Edition', style='CustomStyle').italic = True
    paragraph.insert_paragraph_before().add_run(f'\tTotal Language: {score_data["REELT Total Language Score"]}', style='CustomStyle').bold = True
    paragraph.insert_paragraph_before().add_run(f'\tAuditory Comprehension: {score_data["REELT Auditory Comprehension Score"]}', style='CustomStyle')
    paragraph.insert_paragraph_before().add_run(f'\tExpressive Communication: {score_data["REELT Expressive Communication Score"]}', style='CustomStyle')
    
def add_abas(paragraph, score_data):
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run(f'\t({score_data["Test Date"]}) – Adaptive Behavior Assessment System – Third Edition', style='CustomStyle').italic = True
    paragraph.insert_paragraph_before().add_run(f'\tGeneral Adaptive Composite: {score_data["ABAS General Adaptive Composite"]}', style='CustomStyle').bold = True
    paragraph.insert_paragraph_before().add_run(f'\tConceptual: {score_data["ABAS Conceptual"]}', style='CustomStyle')
    paragraph.insert_paragraph_before().add_run(f'\tSocial: {score_data["ABAS Social"]}\t\t\tPractical: {score_data["ABAS Practical"]}', style='CustomStyle')
    
def add_bullet(paragraph, list_data):
    paragraph.insert_paragraph_before()
    for item in list_data:
        paragraph.insert_paragraph_before().add_run(item, style='ListStyle')
    delete_paragraph(paragraph)


if submit:
    # handle word to replace 
    # pronouns
    with open("misc_data/pronouns.yaml", "r") as file:
        pronoun = yaml.safe_load(file)

    replace_word = {
        "{{Preferred Pronouns 1}}": pronoun[preferred]['pronoun1'],
        "{{Preferred Pronouns 1 CAP}}": pronoun[preferred]['pronoun1cap'],
        "{{Preferred Pronouns 2}}": pronoun[preferred]['pronoun2'],
        "{{Preferred Pronouns 2 CAP}}": pronoun[preferred]['pronoun2cap'],
    }

    replace_word.update(data)

    # Add optional data 
    if not wppsi_score and 'wppsi' in optional:
        del optional['wppsi']
    if not dppr_score and 'dppr' in optional:
        del optional['dppr']
    if not pls_score and 'pls' in optional:
        del optional['pls']
    if not pdms_score and 'pdms' in optional:
        del optional['pdms']
    if not peshv_score and 'peshv' in optional:
        del optional['peshv']
    if not reelt_score and 'reelt' in optional:
        del optional['reelt']
    if not abas_score and 'abas' in optional:
        del optional['abas']

    # Display data 
    yaml_string = yaml.dump(replace_word, sort_keys=False)
    yaml_string = yaml_string + '\n' + yaml.dump(optional, sort_keys=False)
    yaml_data = st.code(yaml_string, language=None)
    

    #### Edit document 
    doc = Document('templates/template_mod_12_noScore.docx')
    if doc:
        # Get file name
        today_date = format_date_with_ordinal(datetime.date.today())
        filename = f"{data['{{Patient First Name}}']} {data['{{Patient Last Name}}']} {today_date}.docx"
        
        ### create document style
        custom_style = doc.styles.add_style('CustomStyle', WD_STYLE_TYPE.CHARACTER)
        custom_style.font.size = Pt(12)
        custom_style.font.name = 'Georgia'

        # list_style = doc.styles.add_style('ListStyle', WD_STYLE.LIST_BULLET)
        # list_style.font.size = Pt(12)
        # list_style.font.name = 'Georgia'

        # Add scores 
        if len(optional) > 0:
            for i, paragraph in enumerate(doc.paragraphs):
                if "Scores are reported here as standard scores" in paragraph.text:
                    if 'wppsi' in optional:
                        add_wppsi(paragraph, optional['wppsi'])
                    if 'dppr' in optional:
                        add_dppr(paragraph, optional["dppr"])
                    if 'pls' in optional:
                        add_pls(paragraph, optional["pls"])
                    if 'pdms' in optional:
                        add_pdms(paragraph, optional["pdms"])
                    if 'peshv' in optional:
                        add_peshv(paragraph, optional['peshv'])
                    if 'reelt' in optional:
                        add_reelt(paragraph, optional['reelt'])
                    if 'abas' in optional:
                        add_abas(paragraph, optional['abas'])
                
                if "SRS Report Information" in paragraph.text:
                    if len(teacher_score) == 0:
                        add_srs_no_teacher(paragraph)
                    else:
                        add_srs_yes_teacher(paragraph, teacher_score)

        # Edit document
        for word in replace_word:
            docxedit.replace_string(doc, old_string=word, new_string=replace_word[word])

        # Replace for lists separated by comma:
        for word in comma:
            new_word = ", ".join(comma[word])
            docxedit.replace_string(doc, old_string=word, new_string=new_word)

        # Replace for lists separated by new line:
        for word in lines:
            new_word = "\n".join(lines[word])
            docxedit.replace_string(doc, old_string=word, new_string=new_word)

        # Save content to file
        doc.save(bio)

        # Replace for lists separated by bullet points
        tpl=DocxTemplate(filename)

        tpl.render(bullet)
        tpl.save(filename)
        
        # Download 
        bio = io.BytesIO()
        document = Document(filename)
        document.save(bio)
        
        st.download_button(
            label="Click here to download",
            data=bio.getvalue(),
            file_name=,
            mime="docx"
        )