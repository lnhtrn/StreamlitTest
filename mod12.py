import streamlit as st
from docx import Document
import docx
import yaml
import io
import docxedit
import datetime
from docx.enum.style import WD_STYLE_TYPE
from streamlit_gsheets import GSheetsConnection
from docxtpl import DocxTemplate
from docx.shared import Inches, Pt
from openai import OpenAI
from modules.recommendations import *

##########################################################
st.set_page_config(
    page_title="Module 1&2",
    page_icon="📝",
    layout="centered",
    initial_sidebar_state="expanded",
)

##########################################################
# Set up OpenAI 
if 'behavior_observation_mod12' not in st.session_state:
    st.session_state.behavior_observation_mod12 = ""
if 'development_history_mod12' not in st.session_state:
    st.session_state.development_history_mod12 = ""

# Load OpenAI client 
client = OpenAI(api_key=st.secrets["openai_key"])

##################################################################
def transcribe_audio(audio_file, name='temp'):
    if audio_file:
        # Transcribe
        with st.spinner("Transcribing...", show_time=True):
            # result = whisper_model.transcribe(f"{name}.wav")
            result = client.audio.transcriptions.create(
                model="whisper-1", 
                file=audio_file, 
                response_format="text"
            )
        return result 

##########################################################
# Access Google Sheets

def get_abbreviation(test_name):
    # Split on en dash (\u2013)
    main_title = test_name.split('\u2013')[0].strip()
    
    # Split into words and get uppercase initials
    abbreviation = ''.join(word[0] for word in main_title.split() if word[0].isupper())
    
    return abbreviation

dropdowns = {}
connections = {}

# Create a connection object.
connections['All'] = st.connection(f"mod12_all", type=GSheetsConnection)

# Read object
df = connections['All'].read(
    ttl="30m",
    usecols=list(range(6)),
    nrows=30,
) 
for col_name in df.columns:
    dropdowns[col_name] = df[col_name].tolist()
    dropdowns[col_name] = [x for x in dropdowns[col_name] if str(x) != 'nan']

# DSM dropdowns
connections['DSM'] = st.connection(f"dsm", type=GSheetsConnection)
# Read object
df = connections['DSM'].read(
    ttl="30m",
    usecols=list(range(7)),
    nrows=15,
) 
for col_name in df.columns:
    dropdowns[col_name] = df[col_name].tolist()
    dropdowns[col_name] = [x for x in dropdowns[col_name] if str(x) != 'nan']
    dropdowns[col_name].append("None")

# Scores for sidebar
connections['Scores'] = st.connection(f"mod12_scores", type=GSheetsConnection)
# Read object
df = connections['Scores'].read(
    ttl="30m",
    usecols=list(range(6)),
    nrows=30,
) 
score_list = df.to_dict('records')

# Process data
scores = {}
check_scores = {}

for test in score_list:
    test_name = test["Test name"]
    abbr = get_abbreviation(test_name)
    scores[abbr] = {}

    scores[abbr]["Test name"] = test_name
    all_lines = []
    all_items = {}
    print(f"\nTest: {test_name}")
    
    for i in range(5):
        line_key = f"Line {i}"
        line_value = test.get(line_key)
        if line_value and str(line_value) != "nan":
            all_lines.append([])
            items = [item.strip() for item in line_value.split(",")]
            for item in items:
                bold = "(bold)" in item
                item_name = item.replace("(bold)", "").strip()
                # write_item(item_name, bold=bold)
                all_lines[i].append((item_name, bold))
                all_items[item_name] = 0
    
    scores[abbr]["Lines"] = all_lines
    scores[abbr]["All items"] = all_items

##################################################
# Set up side bar
def clear_my_cache():
    st.cache_data.clear()

with st.sidebar:
    st.markdown("**After editing dropdown options, please reload data using the button below to update within the form.**")
    st.link_button("Edit Dropdown Options", st.secrets['mod12_spreadsheet'])
    st.link_button("Edit Score Options", st.secrets['mod12_scores'])
    st.button('Reload Dropdown Data', on_click=clear_my_cache)

    # Display data 
    # yaml_dropdown = yaml.dump(dropdowns, sort_keys=False)
    # st.code(yaml_dropdown, language=None)
    
    ####################################################
    st.markdown("**Check to include score in the form:** Scores to report:")
    scq_result = st.checkbox("Social Communication Questionnaire (SCQ) - Lifetime Form")
    teacher_eval = st.checkbox("Teacher's SRS Scores")
    
    for item in scores:
        check_scores[item] = st.checkbox(scores[item]["Test name"])

col1,col2 = st.columns(2)
col1.title('Module 1&2 Report Builder')

def format_date_with_ordinal(date_obj):
    day = date_obj.day
    suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return date_obj.strftime(f"%B {day}{suffix}, %Y")

# Set up dictionary to store data 
data = {}
optional = {}
teacher_score = {}
bullet = {}
lines = {}
comma = {}

# set up recommendation system
check_rec = {}
with open("misc_data/rec_per_module.yaml", "r") as file:
    recommendation_options = yaml.safe_load(file)['mod_12']

#############################################################
# Start of form 
st.header("Appointment Summary")

data['{{Patient First Name}}'] = st.text_input('Patient First Name')

data['{{Patient Last Name}}'] = st.text_input('Patient Last Name')

preferred = st.selectbox(
    "Patient's Preferred Pronoun",
    ("They/them", "He/him", "She/her"),
)

data['{{Location of the evaluation}}'] = st.radio(
    "Location of the evaluation",
    ['home', 'school', 'the office'],
    index=None,
)

# Audio section 
st.markdown(f"**Behavioral Observation:** Things to mention: eye contact, attention to task, social affect and restricted and repetitive behavior.")
audio_behavior = st.audio_input("Behavioral Observation")
if audio_behavior:
    # 3. Create a download button
    st.download_button(
        label="Download Behavioral Observation Recording",
        key="audio_behavior",
        data=audio_behavior,
        file_name=f"{data['{{Patient First Name}}']} {data['{{Patient Last Name}}']} - Behavioral Observation.wav",
        mime="audio/wav",
    )

st.markdown(f"**Developmental History:** Things to mention: social communication skills, repetitive behavior and related behavioral concerns.")
audio_development = st.audio_input("Developmental History")
if audio_development:
    # 3. Create a download button
    st.download_button(
        label="Download Developmental History Recording",
        key="audio_development",
        data=audio_development,
        file_name=f"{data['{{Patient First Name}}']} {data['{{Patient Last Name}}']} - Developmental History.wav",
        mime="audio/wav",
    )

if st.button("Transcribe"):
    if audio_behavior and audio_development:
        transcript_behavior = transcribe_audio(audio_behavior, name='behavior')
        st.markdown(f"**Transcription:** {transcript_behavior}")

        transcript_development = transcribe_audio(audio_development, name='development')
        st.markdown(f"**Transcription:** {transcript_development}")
        
        response = client.responses.create(
            prompt={
                "id": st.secrets["behavior_prompt_mod12_id"],
                # "version": "3",
                "variables": {
                    "first_name": data['{{Patient First Name}}'],
                    "pronouns": preferred,
                    "evaluation_location": data['{{Location of the evaluation}}'],
                    "transcription": transcript_behavior
                }
            }
        )
        st.session_state.behavior_observation_mod12 = response.output_text

        response = client.responses.create(
            prompt={
                "id": st.secrets["development_prompt_mod12_id"],
                # "version": "5",
                "variables": {
                    "first_name": data['{{Patient First Name}}'],
                    "pronouns": preferred,
                    "transcription": transcript_development
                }
            }
        )
        st.session_state.development_history_mod12 = response.output_text
        
################################################################
# Start form
with st.form('BasicInfo'):
    st.header("Patient's data")
    
    data["{{Patient Age}}"] = st.number_input("Patient's Age", 0, 100)
    data['{{Patient age unit}}'] = st.radio(
        "Year/month?",
        ("year", "month")
    )

    data['{{Caregiver type}}'] = st.selectbox(
        "Patient's Caregiver",
        ("mother", "father", "parent", "grandparent", "legal custodian", "foster parent"),
        placeholder="Select from the choices or enter a new one",
        index=None,
        accept_new_options=True,
    )

    bullet['CaregiverPrimaryConcerns'] = st.multiselect(
        "Caregiver\'s Primary Concerns",
        dropdowns["Caregiver\'s Primary Concerns"],
        # [
        #     "Speech delays impacting social opportunities.",
        #     "Clarifying diagnostic presentation.",
        #     "Determining service eligibility.",
        #     "Language delays and difficulties.",
        #     "Elopement and related safety concerns.",
        #     "Determining appropriate supports."
        # ],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )
    
    data['{{Residence City/State}}'] = st.text_input("Residence City/State")
    # st.selectbox(
    #     "Residence City/State", states, index=None,
    # )

    data['{{Narrative}}'] = st.text_area('Narrative to finish \"Patient lives with...\"')

    ##########################################################
    st.header("BRH Evaluation Details")

    data['{{Evaluation Date}}'] = format_date_with_ordinal(st.date_input("Evaluation Date"))

    data['{{Module used}}'] = st.radio("Module used", ["Module 1", "Module 2"])
    if data['{{Module used}}'] == "Module 1":
        data['{{Module Description}}'] = "Module 1 is designed for children with single words"
    else:
        data['{{Module Description}}'] = "Module 2 is designed for children with phrase speech"

    data['{{Results Shared Date}}'] = format_date_with_ordinal(st.date_input("Results Shared Date"))
    
    data['{{Date Report Sent to Patient}}'] = format_date_with_ordinal(st.date_input("Date Report Sent to Patient"))

    lines["{{Result of the evaluation}}"] = st.multiselect(
        "Result of the evaluation",
        dropdowns["Result of the evaluation"],
        # [
        #     "F84.0 - Autism Spectrum Disorder (per the above referenced evaluation)",
        #     "F88.0 - Global Developmental Delay (per behavioral presentation)",
        #     "F80.2 - Mixed Receptive-Expressive Language Disorder",
        #     "F90.2 - Attention Deficit Hyperactivity Disorder - Combined-Type",
        #     "F50.82 Avoidant/Restrictive Food Intake Disorder",
        #     "None"
        # ],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )

    if scq_result:
        data["{{Results (SCQ) - Lifetime Form}}"] = st.text_input(
            "Results (SCQ) - Lifetime Form"
        )

    data["{{SRS-2 Score Caregiver}}"] = st.text_input("Caregiver's SRS-2 Score")
    
    data["{{Social Communication and Interaction Score Caregiver}}"] = st.text_input("Social Communication and Interaction Score Caregiver")
    
    data["{{Restricted Interests and Repetitive Behavior Score Caregiver}}"] = st.text_input("Restricted Interests and Repetitive Behavior Score Caregiver")

    data["{{Caregiver's level of concern}}"] = st.radio(
        "Caregiver's level of concern",
        ['no', 'mild', 'moderate', 'severe']
    )

    data["{{Evaluator's level of concern}}"] = st.radio(
        "Evaluator's level of concern",
        ['no', 'mild', 'moderate', 'severe']
    )

    ##########################################################
    if teacher_eval:
        st.header("Teacher SRS Score")
        st.markdown("*Skip this section if teacher did not give SRS Score*")

        teacher_score['{{Teacher name, title}}'] = st.text_input("Teacher name, title")

        teacher_score['{{SRS-2 Score Teacher}}'] = st.text_input("Teacher's SRS-2 Score")

        teacher_score['{{Social Communication and Interaction Score Teacher}}'] = st.text_input("Social Communication and Interaction Score Teacher")

        teacher_score['{{Restricted Interests and Repetitive Behavior Score Teacher}}'] = st.text_input("Restricted Interests and Repetitive Behavior Score Teacher")

        teacher_score["{{Teacher level of concern}}"] = st.radio(
            "Teacher's level of concern",
            ['no', 'mild', 'moderate', 'severe']
        )

    ######################################################
    st.header("Medical/Developmental History")
    
    lines['{{Diagnosis History}}'] = st.multiselect(
        "Diagnosis History",
        dropdowns['Diagnosis History'],
        # ['History of language and social communication delays.'],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )

    lines['{{Medications}}'] = st.multiselect(
        "Medications",
        ['None noted or reported.'],
        placeholder="Can input multiple options",
        accept_new_options=True
    )

    ###############################################
    st.header("Educational Background")

    data['{{School District}}'] = st.selectbox(
        "School District",
        ['Rochester City'],
        index=None,
        placeholder="Select a school district or enter a new one",
        accept_new_options=True,
    )

    data['{{School Name}}'] = st.text_input("School Name")

    data['{{Grade}}'] = st.text_input(
        "Grade",
        # dropdowns['Grade'],
        # index=None,
        # placeholder="Select a grade or enter a new one",
        # accept_new_options=True,
    )

    data['School Year'] = st.text_input(
        "School Year",
        # dropdowns["School Year"],
        # index=None,
        # placeholder="Select a school year or enter a new one",
        # accept_new_options=True,
    )

    data['{{Education Setting}}'] = st.selectbox(
        "Education Setting",
        ["General Education", "Integrated Co-Taught", "12:1:1", "8:1:1", "6:1:1"],
        index=None,
        placeholder="Select an education setting or enter a new one",
        accept_new_options=True,
    )

    comma['{{Services}}'] = st.multiselect(
        "Services",
        dropdowns['Services'],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )

    ##########################################################
    # Score section
    for test in check_scores:
        if check_scores[test]:
            st.header(scores[test]["Test name"])
            st.markdown(f"*Skip this section if there is no {test} Score*")

            scores[test]["Test Date"] = st.date_input(f"{test} Test Date").strftime("%m/%Y")
            
            for item in scores[test]["All items"]:
                scores[test]["All items"][item] = st.text_input(item)    

    ########################################################
    st.header("Behavioral Presentation")
    data['behavior_observation'] = st.text_area(
        "Behavioral Observation: Edit the response before submitting the form", 
        # behavior_observation,
        st.session_state.behavior_observation_mod12,
        height=400,
    )

    ########################################################
    st.header("Developmental History")
    data['development_history'] = st.text_area(
        "Developmental History: Edit the response before submitting the form", 
        # development_history,
        st.session_state.development_history_mod12,
        height=400,
    )

    ########################################################
    st.header("DSM Criteria")
    # Add DSM questions
    with open("misc_data/dsm.yaml", "r") as file:
        dsm_title = yaml.safe_load(file)
    
    for key, value in dsm_title.items():
        bullet[key] = st.multiselect(    
            value,
            dropdowns[key],
            placeholder="Select multiple options from the list or enter a new one",
            accept_new_options=True
        )

    comma['{{Symptoms present in the early developmental period}}'] = st.multiselect(
        "Symptoms present in the early developmental period",
        [
            "Confirmed by record review",
            "None",
        ],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )

    comma['{{Symptoms cause clinically significant impairment}}'] = st.multiselect(
        "Symptoms cause clinically significant impairment",
        [
            "Confirmed by record review",
            "None",
        ],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )

    ########################################################################
    st.header("Recommendations")

    for key, label in recommendation_options.items():
        check_rec[key] = st.checkbox(label)
    
    # data['{{}}'] = st.text_input("")
    # data['{{}}'] = st.text_input("")
    # data['{{}}'] = st.text_input("")

    submit = st.form_submit_button('Submit')


def add_behavior_presentation(paragraph, transcript):
    # separate transcript
    small_para = transcript.split('\n\n')

    st.write(small_para)

    paragraph.insert_paragraph_before().add_run(small_para[0], style='CustomStyle')
    paragraph.insert_paragraph_before()

    for sub_para in small_para[1:]:
        sub_para = sub_para.split(":")
        p = paragraph.insert_paragraph_before()
        p.add_run(sub_para[0], style='CustomStyle').italic = True
        p.add_run(f":{sub_para[1]}\n", style='CustomStyle')
        
    delete_paragraph(paragraph)

def add_developmental_history(paragraph, transcript):
    # separate transcript
    small_para = transcript.split('\n\n')
    st.write(small_para)

    for sub_para in small_para:
        sub_para = sub_para.split(":")
        p = paragraph.insert_paragraph_before()
        p.add_run(sub_para[0], style='CustomStyle').italic = True
        p.add_run(f":{sub_para[1]}\n", style='CustomStyle')
        
    delete_paragraph(paragraph)

def add_school(paragraph):
    p = paragraph.insert_paragraph_before()
    tab_stops = p.paragraph_format.tab_stops
    # tab_stops.clear()  # Start fresh for this paragraph only
    tab_stops.add_tab_stop(Inches(3))
    # Add data
    p.add_run("District", style='CustomStyle').font.underline = True
    p.add_run(f": {data['{{School District}}']}\t", style='CustomStyle')
    p.add_run("Grade", style='CustomStyle').font.underline = True
    ### italics for school year
    p.add_run(f": {data['{{Grade}}']} (", style='CustomStyle')
    p.add_run({data['School Year']}, style='CustomStyle').italic = True
    p.add_run(")\n\n", style='CustomStyle')
    p.add_run("School", style='CustomStyle').font.underline = True
    p.add_run(f": {data['{{School Name}}']}\t", style='CustomStyle')
    p.add_run("Setting", style='CustomStyle').font.underline = True
    p.add_run(f": {data['{{Education Setting}}']}", style='CustomStyle')
    delete_paragraph(paragraph)

def add_scq_form(paragraph):
    r = paragraph.insert_paragraph_before().add_run('Social Communication Questionnaire (SCQ) – Lifetime Form', style='CustomStyle')
    r.italic = True
    r.font.underline = True
    p = paragraph.insert_paragraph_before()
    p.add_run("The SCQ evaluates for symptoms of autism spectrum disorder across developmental history. Scores above 15 are suggestive of an autism diagnosis. Based on {{Preferred Pronouns 2}} {{Caregiver type}}’s report, {{Patient First Name}}’s score was {{Results (SCQ) - Lifetime Form}}. ", style='CustomStyle')
    p.add_run("This score is clearly consistent with autism at present.\n", style='CustomStyle').italic = True

def add_srs_no_teacher(paragraph):
    r = paragraph.insert_paragraph_before().add_run('Social Responsiveness Scale – Second Edition (SRS-2) – Parent Report', style='CustomStyle')
    r.italic = True
    r.font.underline = True
    paragraph.insert_paragraph_before().add_run('The SRS-2 is an objective measure that identifies social impairments associated with autism spectrum disorder and quantifies ASD-related severity throughout the lifespan. \nThe following interpretative guidelines are offered here for the benefit of the reader: Less than 59 indicates within normal limits, between 60 and 65 as mild concern, between 65 and 75 as moderate concern, and greater than 76 as severe concern. ', style='CustomStyle')
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run('\tSRS-2 Total Score: {{SRS-2 Score Caregiver}} ({{Caregiver type}})', style='CustomStyle').bold = True
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run('\tSocial Communication and Interaction: {{Social Communication and Interaction Score Caregiver}} ({{Caregiver type}})', style='CustomStyle')
    paragraph.insert_paragraph_before().add_run('\tRestricted Interests and Repetitive Behavior: {{Restricted Interests and Repetitive Behavior Score Caregiver}} ({{Caregiver type}})', style='CustomStyle')
    paragraph.insert_paragraph_before()
    observe = paragraph.insert_paragraph_before()
    observe.add_run("Based on the report provided by {{Preferred Pronouns 2}} {{Caregiver type}}, ", style='CustomStyle')
    observe.add_run("{{Patient First Name}}’s social communication and related behaviors indicated {{Caregiver's level of concern}} concerns. ", style='CustomStyle').italic = True
    observe.add_run("My observation aligned with a {{Evaluator's level of concern}} concern.", style='CustomStyle').bold = True
    delete_paragraph(paragraph)

def add_srs_yes_teacher(paragraph, score_data):
    r = paragraph.insert_paragraph_before().add_run('Social Responsiveness Scale – Second Edition (SRS-2) – Parent & Teacher Report', style='CustomStyle')
    r.italic = True
    r.font.underline = True
    paragraph.insert_paragraph_before().add_run('The SRS-2 is an objective measure that identifies social impairments associated with autism spectrum disorder and quantifies ASD-related severity throughout the lifespan. \nThe following interpretative guidelines are offered here for the benefit of the reader: Less than 59 indicates within normal limits, between 60 and 65 as mild concern, between 65 and 75 as moderate concern, and greater than 76 as severe concern. ', style='CustomStyle')
    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before()
    p.add_run('\tSRS-2 Total Score: {{SRS-2 Score Caregiver}} ({{Caregiver type}}), ', style='CustomStyle').bold = True
    p.add_run(f"{score_data['{{SRS-2 Score Teacher}}']} (teacher)", style='CustomStyle').bold = True
    paragraph.insert_paragraph_before()
    p = paragraph.insert_paragraph_before()
    p.add_run('\tSocial Communication and Interaction: {{Social Communication and Interaction Score Caregiver}} ({{Caregiver type}}), ', style='CustomStyle')
    p.add_run(f"{score_data['{{Social Communication and Interaction Score Teacher}}']} (teacher)", style='CustomStyle')
    p = paragraph.insert_paragraph_before()
    p.add_run('\tRestricted Interests and Repetitive Behavior: {{Restricted Interests and Repetitive Behavior Score Caregiver}} ({{Caregiver type}}), ', style='CustomStyle')
    p.add_run(f'{score_data["{{Restricted Interests and Repetitive Behavior Score Teacher}}"]} (teacher)', style='CustomStyle')
    paragraph.insert_paragraph_before()
    observe = paragraph.insert_paragraph_before()
    observe.add_run("Based on the report provided by {{Preferred Pronouns 2}} {{Caregiver type}}, ", style='CustomStyle')
    observe.add_run("{{Patient First Name}}’s social communication and related behaviors indicated {{Caregiver's level of concern}} concerns. ", style='CustomStyle').italic = True
    observe.add_run("{{Patient First Name}}’s teacher reported a ", style='CustomStyle')
    observe.add_run(f"{score_data['{{Teacher level of concern}}']} concern, and ", style='CustomStyle')
    observe.add_run("my observation aligned with a {{Evaluator's level of concern}} concern.", style='CustomStyle').bold = True
    delete_paragraph(paragraph)

def add_score(paragraph, score_data):
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run(f'\t({score_data["Test Date"]}) \u2013 {score_data["Test name"]}', style='CustomStyle').italic = True
    
    # Go over each line 
    for line in score_data["Lines"]:
        # get a new paragraph and indent it 
        p = paragraph.insert_paragraph_before()
        p.paragraph_format.left_indent = Inches(0.5)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(3.5))

        # add each score
        for item_tuple in line:
            item = item_tuple[0]
            p.add_run(f'{item}: {score_data["All items"][item]}\t', style='CustomStyle').bold = item_tuple[1]

if submit:
    # Update session state 
    st.session_state.behavior_observation_mod12 = data['behavior_observation']
    st.session_state.development_history_mod12 = data['development_history'] 

    # handle word to replace 
    # pronouns
    with open("misc_data/pronouns.yaml", "r") as file:
        pronoun = yaml.safe_load(file)

    replace_word = {
        "{{Preferred Pronouns 1}}": pronoun[preferred]['pronoun1'],
        "{{Preferred Pronouns 1 CAP}}": pronoun[preferred]['pronoun1cap'],
        "{{Preferred Pronouns 2}}": pronoun[preferred]['pronoun2'],
        "{{Preferred Pronouns 2 CAP}}": pronoun[preferred]['pronoun2cap'],
        "{{Gender}}": pronoun[preferred]['gender'],
    }

    replace_word.update(data)

    # Display data 
    yaml_string = yaml.dump(replace_word, sort_keys=False)
    yaml_string = yaml_string + '\n' + yaml.dump(optional, sort_keys=False)
    yaml_string = yaml_string + '\n' + yaml.dump(bullet, sort_keys=False)
    yaml_data = st.code(yaml_string, language=None)
    

    #### Edit document 
    doc = Document('templates/template_mod_12.docx')
    if doc:
        # Get file name
        today_date = format_date_with_ordinal(datetime.date.today())
        filename = f"{data['{{Patient First Name}}']} {data['{{Patient Last Name}}']} {today_date}.docx"
        
        ### create document 
        norm_style = doc.styles['Normal']
        norm_style.paragraph_format.line_spacing = 1

        custom_style = doc.styles.add_style('CustomStyle', WD_STYLE_TYPE.CHARACTER)
        custom_style.font.size = Pt(12)
        custom_style.font.name = 'Georgia'

        custom_style_2 = doc.styles.add_style('CustomStyle2', WD_STYLE_TYPE.CHARACTER)
        custom_style_2.font.size = Pt(11.5)
        custom_style_2.font.name = 'Georgia'

        list_style = doc.styles['Bullet New']
        list_style.paragraph_format.line_spacing = 1

        # Add scores 
        for i, paragraph in enumerate(doc.paragraphs):
            if "Scores are reported here as standard scores" in paragraph.text:
                total = 0
                for test in check_scores:
                    if check_scores[test]:
                        total += 1 
                        if total == 1:
                            paragraph.insert_paragraph_before().add_run("Psychoeducational Testing:\n", style='CustomStyle').font.underline = True
                        add_score(paragraph, score_data=scores[test])

                if total == 0:
                    delete_paragraph(paragraph)

            if "[[Behavioral Presentation]]" in paragraph.text:
                add_behavior_presentation(paragraph, st.session_state.behavior_observation_mod12)
            
            if "[[Developmental History]]" in paragraph.text:
                add_developmental_history(paragraph, st.session_state.development_history_mod12)

            if "[[Recommendations]]" in paragraph.text:
                for rec, checked in check_rec.items():
                    if checked:
                        func = globals().get(f"add_{rec}")
                        if callable(func):
                            func(paragraph)
                
                delete_paragraph(paragraph)
                
            if "SRS Report Information" in paragraph.text:
                # Add SCQ
                if scq_result:
                    add_scq_form(paragraph)
                # Add SRS
                if len(teacher_score) == 0:
                    add_srs_no_teacher(paragraph)
                else:
                    add_srs_yes_teacher(paragraph, teacher_score)
            
            if "Social Responsiveness Scale" in paragraph.text:
                if teacher_eval:
                    paragraph.add_run(" & teacher\nDevelopmental History & Review of Records\n", style='CustomStyle')
                    paragraph.add_run(f"School Report on SRS-2 provided by {teacher_score['{{Teacher name, title}}']}", style='CustomStyle')
                else:
                    paragraph.add_run("\nDevelopmental History & Review of Records", style='CustomStyle')

            if "[[District Grade School Setting]]" in paragraph.text:
                add_school(paragraph)
        
        # Edit document
        for word in replace_word:
            docxedit.replace_string(doc, old_string=word, new_string=replace_word[word])

        # Replace for lists separated by comma:
        for word in comma:
            new_word = ", ".join(comma[word])
            docxedit.replace_string(doc, old_string=word, new_string=new_word)

        # Replace for lists separated by new line:
        for word in lines:
            new_word = "\n".join(lines[word])
            docxedit.replace_string(doc, old_string=word, new_string=new_word)

        # Save content to file
        doc.save(filename)

        # Replace for lists separated by bullet points
        tpl=DocxTemplate(filename)
        print("Load template!")

        tpl.render(bullet)
        print("Bullet rendered!")

        tpl.save(filename)
        print("File saved at", filename)

        # Download 
        bio = io.BytesIO()
        document = Document(filename)
        document.save(bio)
        
        st.download_button(
            label="Click here to download",
            key="report_download",
            data=bio.getvalue(),
            file_name=filename,
            mime="docx"
        )