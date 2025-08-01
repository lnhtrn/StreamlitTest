import streamlit as st
from docx import Document
import docx
import yaml
import io
import docxedit
import datetime
from docx.enum.style import WD_STYLE_TYPE
from streamlit_gsheets import GSheetsConnection
from docxtpl import DocxTemplate
from docx.shared import Inches, Pt
from openai import OpenAI
from modules.recommendations import *

##########################################################
st.set_page_config(
    page_title="Recommendation Testing",
    page_icon="📝",
    layout="centered",
    initial_sidebar_state="expanded",
)

# Set up authentication 
if not st.user.is_logged_in:
    col1, col2, col3 = st.columns([1, 2, 1]) # Adjust ratios as needed for desired centering
    with col2:
        st.title("Log in to use Report Builder!")
        if st.button("Log in with Google Account"):
            st.login("google")
    st.stop()

# Sidebar after logging in 
st.sidebar.write(f"Welcome, {st.user.name}!")
if st.sidebar.button("Log out"):
    st.logout()

##########################################################
# Access Google Sheets

def get_abbreviation(test_name):
    # Split on en dash (\u2013)
    main_title = test_name.split('\u2013')[0].strip()
    
    # Split into words and get uppercase initials
    abbreviation = ''.join(word[0] for word in main_title.split() if word[0].isupper())
    
    return abbreviation

dropdowns = {}
connections = {}

# Create a connection object.
connections['All'] = st.connection(f"mod12_all", type=GSheetsConnection)

# Read object
df = connections['All'].read(
    ttl="30m",
    usecols=list(range(6)),
    nrows=30,
) 
for col_name in df.columns:
    dropdowns[col_name] = df[col_name].tolist()
    dropdowns[col_name] = [x for x in dropdowns[col_name] if str(x) != 'nan']

# DSM dropdowns
connections['DSM'] = st.connection(f"dsm", type=GSheetsConnection)
# Read object
df = connections['DSM'].read(
    ttl="30m",
    usecols=list(range(7)),
    nrows=15,
) 
for col_name in df.columns:
    dropdowns[col_name] = df[col_name].tolist()
    dropdowns[col_name] = [x for x in dropdowns[col_name] if str(x) != 'nan']
    dropdowns[col_name].append("None")

# Scores for sidebar
connections['Scores'] = st.connection(f"mod12_scores", type=GSheetsConnection)
# Read object
df = connections['Scores'].read(
    ttl="30m",
    usecols=list(range(6)),
    nrows=30,
) 
score_list = df.to_dict('records')

# Process data
scores = {}
check_scores = {}

for test in score_list:
    test_name = test["Test name"]
    abbr = get_abbreviation(test_name)
    scores[abbr] = {}

    scores[abbr]["Test name"] = test_name
    all_lines = []
    all_items = {}
    print(f"\nTest: {test_name}")
    
    for i in range(5):
        line_key = f"Line {i}"
        line_value = test.get(line_key)
        if line_value and str(line_value) != "nan":
            all_lines.append([])
            items = [item.strip() for item in line_value.split(",")]
            for item in items:
                bold = "(bold)" in item
                item_name = item.replace("(bold)", "").strip()
                # write_item(item_name, bold=bold)
                all_lines[i].append((item_name, bold))
                all_items[item_name] = 0
    
    scores[abbr]["Lines"] = all_lines
    scores[abbr]["All items"] = all_items

################ RECOMMENDATION #################
rec_dict = {}

# Connect Google Sheets for Recommendation 
connections['Recommendation'] = st.connection(f"recommendations", type=GSheetsConnection)

# Read object
df = connections['Recommendation'].read(
    ttl="30m",
    usecols=list(range(2)),
    nrows=200,
) 

for _, row in df.iterrows():
    key = row['Title']
    values = []
    for para in row['Content'].split('\n'):
        para_value = []
        for item in para.split(';'):
            item = item.strip()
            if '[' in item and ']' in item:
                data_part = item.split('[')[0].strip()
                format_part = item.split('[')[1].replace(']', '').strip()
                para_value.append((data_part, format_part))
        values.append(para_value)
    rec_dict[key] = values

connections['Recommendation_Per_Module'] = st.connection(f"recommendations_per_module", type=GSheetsConnection)

# Read object
df = connections['Recommendation_Per_Module'].read(
    ttl="30m",
    usecols=list(range(2)),
    nrows=200,
) 

rec_list = df[df["Module Name"] == 'Module 4']["Recommendation Name"].tolist()

##################################################
# Set up side bar
def clear_my_cache():
    st.cache_data.clear()

with st.sidebar:
    st.markdown("**After editing dropdown options, please reload data using the button below to update within the form.**")
    st.link_button("Edit Dropdown Options", st.secrets['mod12_spreadsheet'])
    st.link_button("Edit Score Options", st.secrets['mod12_scores'])
    st.button('Reload Dropdown Data', on_click=clear_my_cache)

    # Display data 
    # yaml_dropdown = yaml.dump(dropdowns, sort_keys=False)
    # st.code(yaml_dropdown, language=None)
    
    ####################################################
    # st.markdown("**Check to include score in the form:** Scores to report:")
    # scq_result = st.checkbox("Social Communication Questionnaire (SCQ) - Lifetime Form")
    # teacher_eval = st.checkbox("Teacher's SRS Scores")
    
    # for item in scores:
    #     check_scores[item] = st.checkbox(scores[item]["Test name"])

# col1,col2 = st.columns(2)
# col1
st.title('Recommendation Testing')
st.markdown("*For authorized use by Bryan R. Harrison, PhD Psychologist, PC only.*")
st.markdown("---")

def format_date_with_ordinal(date_obj):
    day = date_obj.day
    suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return date_obj.strftime(f"%B {day}{suffix}, %Y")

# Set up dictionary to store data 
data = {}

# set up recommendation system
check_rec = {}
with open("misc_data/rec_per_module.yaml", "r") as file:
    recommendation_options = yaml.safe_load(file)['mod_12']

# Display data 
# yaml_string = yaml.dump(rec_dict, sort_keys=False)
# yaml_data = st.code(yaml_string, language=None)

#############################################################
# Start of form 
st.header("Appointment Summary")

data['{{Patient First Name}}'] = st.text_input('Patient First Name')

data['{{Patient Last Name}}'] = st.text_input('Patient Last Name')

preferred = st.selectbox(
    "Patient's Preferred Pronoun",
    ("They/them", "He/him", "She/her"),
)

data['{{Location of the evaluation}}'] = st.radio(
    "Location of the evaluation",
    ['home', 'school', 'the office'],
    index=None,
)

#########################################################
with st.form("BasicInfo"):  
    st.header("Recommendations")

    check_rec = {}
    for rec in rec_list:
        check_rec[rec] = st.checkbox(rec)

    #############################################
    submit = st.form_submit_button('Submit')


###########################################################

if submit:
    replace_word = {}
    replace_percent = {}

    # Edit document 
    doc = Document('templates/template_mod_4.docx')
    if doc:
        ### create document style
        doc_style = doc.styles['Normal']
        font = doc_style.font
        font.name = 'Georgia'
        font.size = Pt(12)

        custom_style = doc.styles.add_style('CustomStyle', WD_STYLE_TYPE.CHARACTER)
        custom_style.font.size = Pt(12)
        custom_style.font.name = 'Georgia'

        
        for paragraph in doc.paragraphs:
            if "[[Recommendations]]" in paragraph.text:
                for rec in check_rec:
                    if check_rec[rec]:
                        rec_item = rec_dict[rec]
                        for para in rec_item:
                            p = paragraph.insert_paragraph_before()
                            for para_item in para:
                                if para_item[1] == "bold":
                                    add_bold(p, para_item[0])
                                elif para_item[1] == "normal":
                                    add_normal(p, para_item[0])
                                elif para_item[1] == "bullet":
                                    add_bullet(p, para_item[0])
                                elif para_item[1] == "link":
                                    add_hyperlink(p, para_item[0])
                                else:
                                    pass
                        paragraph.insert_paragraph_before()
                delete_paragraph(paragraph)

        # Get file name
        filename = f"{data['{{Patient First Name}}']} {data['{{Patient Last Name}}']} test recommendation.docx"

        # Save content to file
        doc.save(filename)

        # Download 
        bio = io.BytesIO()
        document = Document(filename)
        document.save(bio)
        
        st.download_button(
            label="Click here to download",
            key="report_download",
            data=bio.getvalue(),
            file_name=filename,
            mime="docx"
        )


    
    
