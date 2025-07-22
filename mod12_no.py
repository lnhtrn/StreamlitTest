import streamlit as st
from docx import Document
import docx
import yaml
import io
import docxedit
import datetime
from docx.enum.style import WD_STYLE_TYPE
from streamlit_gsheets import GSheetsConnection
from docxtpl import DocxTemplate
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement, qn
from openai import OpenAI
from modules.recommendations import *

##########################################################
st.set_page_config(
    page_title="Module 1&2 No Autism",
    page_icon="📝",
    layout="centered",
    initial_sidebar_state="expanded",
)

##########################################################
# Set up OpenAI 
if 'behavior_observation_mod12_no_autism' not in st.session_state:
    st.session_state.behavior_observation_mod12_no_autism = ""

# Load OpenAI client 
client = OpenAI(api_key=st.secrets["openai_key"])

##################################################################
def transcribe_audio(audio_file, name='temp'):
    if audio_file:
        # Transcribe
        with st.spinner("Transcribing...", show_time=True):
            # result = whisper_model.transcribe(f"{name}.wav")
            result = client.audio.transcriptions.create(
                model="whisper-1", 
                file=audio_file, 
                response_format="text"
            )
        return result 

##########################################################
# Access Google Sheets

def get_abbreviation(test_name):
    # Split on en dash (\u2013)
    main_title = test_name.split('\u2013')[0].strip()
    
    # Split into words and get uppercase initials
    abbreviation = ''.join(word[0] for word in main_title.split() if word[0].isupper())
    
    return abbreviation

dropdowns = {}
connections = {}

# Create a connection object.
connections['All'] = st.connection(f"mod12_all", type=GSheetsConnection)
# Read object
df = connections['All'].read(
    ttl="30m",
    usecols=list(range(6)),
    nrows=30,
) 
for col_name in df.columns:
    dropdowns[col_name] = df[col_name].tolist()
    dropdowns[col_name] = [x for x in dropdowns[col_name] if str(x) != 'nan']

# Create a connection object.
connections['noAutism'] = st.connection(f"mod12_noAutism", type=GSheetsConnection)
# Read object
df = connections['noAutism'].read(
    ttl="30m",
    usecols=list(range(3)),
    nrows=30,
) 
for col_name in df.columns:
    dropdowns[col_name] = df[col_name].tolist()
    dropdowns[col_name] = [x for x in dropdowns[col_name] if str(x) != 'nan']

# Scores for sidebar
connections['Scores'] = st.connection(f"mod12_scores", type=GSheetsConnection)
# Read object
df = connections['Scores'].read(
    ttl="30m",
    usecols=list(range(6)),
    nrows=30,
) 
score_list = df.to_dict('records')

# Process data
scores = {}
check_scores = {}

for test in score_list:
    test_name = test["Test name"]
    abbr = get_abbreviation(test_name)
    scores[abbr] = {}

    scores[abbr]["Test name"] = test_name
    all_lines = []
    all_items = {}
    print(f"\nTest: {test_name}")
    
    for i in range(5):
        line_key = f"Line {i}"
        line_value = test.get(line_key)
        if line_value and str(line_value) != "nan":
            all_lines.append([])
            items = [item.strip() for item in line_value.split(",")]
            for item in items:
                bold = "(bold)" in item
                item_name = item.replace("(bold)", "").strip()
                # write_item(item_name, bold=bold)
                all_lines[i].append((item_name, bold))
                all_items[item_name] = 0
    
    scores[abbr]["Lines"] = all_lines
    scores[abbr]["All items"] = all_items

##################################################
# Set up side bar
def clear_my_cache():
    st.cache_data.clear()

with st.sidebar:
    st.markdown("**After editing dropdown options, please reload data using the button below to update within the form.**")
    st.link_button("Edit Dropdown Options", st.secrets['mod12_spreadsheet'])
    st.link_button("Edit Score Options", st.secrets['mod12_scores'])
    st.button('Reload Dropdown Data', on_click=clear_my_cache)

    # Display data 
    # yaml_dropdown = yaml.dump(dropdowns, sort_keys=False)
    # st.code(yaml_dropdown, language=None)
    
    ####################################################
    st.markdown("**Check to include score in the form:** Scores to report:")
    # scq_result = st.checkbox("Social Communication Questionnaire (SCQ) - Lifetime Form")
    # teacher_eval = st.checkbox("Teacher's SSR Scores")
    for item in scores:
        check_scores[item] = st.checkbox(scores[item]["Test name"])


col1,col2 = st.columns(2)
col1.title('Module 1&2 No Autism Report Builder')

def format_date_with_ordinal(date_obj):
    day = date_obj.day
    suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return date_obj.strftime(f"%B {day}{suffix}, %Y")

# Set up dictionary to store data 
data = {}
teacher_score = {}
bullet = {}
lines = {}
comma = {}

# set up recommendation system
check_rec = {}
with open("misc_data/rec_per_module.yaml", "r") as file:
    recommendation_options = yaml.safe_load(file)['mod_12_no_autism']

####################################################
st.header("Appointment Summary")

data['{{Patient First Name}}'] = st.text_input('Patient First Name')

data['{{Patient Last Name}}'] = st.text_input('Patient Last Name')

preferred = st.selectbox(
    "Patient's Preferred Pronoun",
    ("They/them", "He/him", "She/her"),
)

data['{{Location of the evaluation}}'] = st.radio(
    "Location of the evaluation",
    ['home', 'school', 'the office'],
    index=None,
)

# Audio section 
st.markdown(f"**Behavioral Observation:** Things to mention: eye contact, attention to task, social affect and restricted and repetitive behavior.")
audio_behavior = st.audio_input("Behavioral Observation")
if audio_behavior:
    # 3. Create a download button
    st.download_button(
        label="Download Behavioral Observation Recording",
        key="audio_behavior",
        data=audio_behavior,
        file_name=f"{data['{{Patient First Name}}']} {data['{{Patient Last Name}}']} - Behavioral Observation.wav",
        mime="audio/wav",
    )

if st.button("Transcribe"):
    if audio_behavior:
        transcript_behavior = transcribe_audio(audio_behavior, name='behavior')
        st.markdown(f"**Transcription:** {transcript_behavior}")
        
        response = client.responses.create(
            prompt={
                "id": st.secrets["behavior_prompt_mod12no_id"],
                # "version": "3",
                "variables": {
                    "first_name": data['{{Patient First Name}}'],
                    "pronouns": preferred,
                    "evaluation_location": data['{{Location of the evaluation}}'],
                    "transcription": transcript_behavior
                }
            }
        )
        st.session_state.behavior_observation_mod12_no_autism = response.output_text

####################################################
with st.form('BasicInfo'):
    st.header("Patient's data")

    data["{{Patient Age}}"] = st.number_input("Patient's Age", 0, 100)
    data['{{Patient age unit}}'] = st.radio(
        "Year/month?",
        ("year", "month")
    )

    data['{{Caregiver type}}'] = st.selectbox(
        "Patient's Caregiver",
        ("mother", "father", "parent", "grandparent", "legal custodian", "foster parent"),
        placeholder="Select from the choices or enter a new one",
        index=None,
        accept_new_options=True,
    )

    bullet['CaregiverPrimaryConcerns'] = st.multiselect(
        "Caregiver\'s Primary Concerns",
        dropdowns['Caregiver\'s Primary Concerns'],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )
    
    data['{{Residence City/State}}'] = st.text_input("Residence City/State")

    data['{{Narrative}}'] = st.text_area('Narrative to finish \"Patient lives with...\"')

    ##########################################################
    st.header("BRH Evaluation Details")

    data['{{Evaluation Date}}'] = format_date_with_ordinal(st.date_input("Evaluation Date"))

    data['{{Module used}}'] = st.radio("Module used", ["Module 1", "Module 2"])
    if data['{{Module used}}'] == "Module 1":
        data['{{Module Description}}'] = "Module 1 is designed for children with single words"
    else:
        data['{{Module Description}}'] = "Module 2 is designed for children with phrase speech"

    data['{{Results Shared Date}}'] = format_date_with_ordinal(st.date_input("Results Shared Date"))
    
    data['{{Date Report Sent to Patient}}'] = format_date_with_ordinal(st.date_input("Date Report Sent to Patient"))

    ######################################################
    st.header("Medical/Developmental History")
    
    lines['{{Developmental Concerns}}'] = st.multiselect(
        "Developmental Concerns",
        dropdowns['Developmental Concerns'],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )

    lines['{{Medical Concerns}}'] = st.multiselect(
        "Medical Concerns",
        dropdowns['Medical Concerns'],
        placeholder="Can input multiple options",
        accept_new_options=True
    )

    bullet['CaregiverDevelopmentalConcerns'] = st.multiselect(
        "Caregiver\'s Developmental Concerns",
        dropdowns['Caregiver\'s Developmental Concerns'],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )

    ###############################################
    st.header("Educational Background")

    data['{{School District}}'] = st.selectbox(
        "School District",
        ['Rochester City'],
        index=None,
        placeholder="Select a school district or enter a new one",
        accept_new_options=True,
    )

    data['{{School Name}}'] = st.text_input("School Name")

    data['{{Grade}}'] = st.text_input(
        "Grade",
    )

    data['School Year'] = st.text_input(
        "School Year",
    )

    data['{{Education Setting}}'] = st.selectbox(
        "Education Setting",
        [
            "General Education", 
            "Integrated Co-Taught", 
            "12:1:1", 
            "8:1:1", 
            "6:1:1"
        ],
        index=None,
        placeholder="Select a grade or enter a new one",
        accept_new_options=True,
    )

    comma['{{Services}}'] = st.multiselect(
        "Services",
        dropdowns['Services'],
        placeholder="Select multiple options from the list or enter a new one",
        accept_new_options=True
    )

    ##########################################################
    # Score section
    for test in check_scores:
        if check_scores[test]:
            st.header(scores[test]["Test name"])
            st.markdown(f"*Skip this section if there is no {test} Score*")

            scores[test]["Test Date"] = st.date_input(f"{test} Test Date").strftime("%m/%Y")
            
            for item in scores[test]["All items"]:
                scores[test]["All items"][item] = st.text_input(item)

    ########################################################
    st.header("Behavioral Presentation")
    data['behavior_observation'] = st.text_area(
        "Behavioral Observation: Edit the response before submitting the form", 
        # behavior_observation,
        st.session_state.behavior_observation_mod12_no_autism,
        height=400,
    )

    ########################################################################
    st.header("Recommendations")

    for key, label in recommendation_options.items():
        check_rec[key] = st.checkbox(label)
    
    # data['{{}}'] = st.text_input("")
    # data['{{}}'] = st.text_input("")
    # data['{{}}'] = st.text_input("")

    submit = st.form_submit_button('Submit')


def add_behavior_presentation(paragraph, transcript):
    # separate transcript
    small_para = transcript.split('\n\n')

    st.write(small_para)

    paragraph.insert_paragraph_before().add_run(small_para[0], style='CustomStyle')
    paragraph.insert_paragraph_before()

    for sub_para in small_para[1:]:
        sub_para = sub_para.split(":")
        p = paragraph.insert_paragraph_before()
        p.add_run(sub_para[0], style='CustomStyle').italic = True
        p.add_run(f":{sub_para[1]}\n", style='CustomStyle')
        
    delete_paragraph(paragraph)

def add_school(paragraph):
    p = paragraph.insert_paragraph_before()
    tab_stops = p.paragraph_format.tab_stops
    # tab_stops.clear()  # Start fresh for this paragraph only
    tab_stops.add_tab_stop(Inches(3))
    # Add data
    p.add_run("District", style='CustomStyle').font.underline = True
    p.add_run(f": {data['{{School District}}']}\t", style='CustomStyle')
    p.add_run("Grade", style='CustomStyle').font.underline = True
    ### italics for school year
    p.add_run(f": {data['{{Grade}}']} (", style='CustomStyle')
    p.add_run(f"{data['School Year']})\n\n", style='CustomStyle').italic = True
    p.add_run("School", style='CustomStyle').font.underline = True
    p.add_run(f": {data['{{School Name}}']}\t", style='CustomStyle')
    p.add_run("Setting", style='CustomStyle').font.underline = True
    p.add_run(f": {data['{{Education Setting}}']}", style='CustomStyle')
    delete_paragraph(paragraph)

def add_score(paragraph, score_data):
    paragraph.insert_paragraph_before()
    paragraph.insert_paragraph_before().add_run(f'\t({score_data["Test Date"]}) \u2013 {score_data["Test name"]}', style='CustomStyle').italic = True
    
    # Go over each line 
    for line in score_data["Lines"]:
        # get a new paragraph and indent it 
        p = paragraph.insert_paragraph_before()
        p.paragraph_format.left_indent = Inches(0.5)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(3.5))

        # add each score
        for item_tuple in line:
            item = item_tuple[0]
            p.add_run(f'{item}: {score_data["All items"][item]}\t', style='CustomStyle').bold = item_tuple[1]

if submit:
    # Update session state 
    st.session_state.behavior_observation_mod12_no_autism = data['behavior_observation']

    # handle word to replace 
    # pronouns
    with open("misc_data/pronouns.yaml", "r") as file:
        pronoun = yaml.safe_load(file)

    replace_word = {
        "{{Preferred Pronouns 1}}": pronoun[preferred]['pronoun1'],
        "{{Preferred Pronouns 1 CAP}}": pronoun[preferred]['pronoun1cap'],
        "{{Preferred Pronouns 2}}": pronoun[preferred]['pronoun2'],
        "{{Preferred Pronouns 2 CAP}}": pronoun[preferred]['pronoun2cap'],
    }

    replace_word.update(data)

    # Display data 
    # yaml_string = yaml.dump(replace_word, sort_keys=False)
    # yaml_string = yaml_string + '\n' + yaml.dump(scores, sort_keys=False)
    # yaml_string = yaml_string + '\n' + yaml.dump(bullet, sort_keys=False)
    # yaml_data = st.code(yaml_string, language=None)
    
    #### Edit document 
    doc = Document('templates/template_mod_12_no_autism.docx')
    if doc:
        # Get file name
        today_date = format_date_with_ordinal(datetime.date.today())
        filename = f"{data['{{Patient First Name}}']} {data['{{Patient Last Name}}']} {today_date}.docx"
        
        ### create document 
        norm_style = doc.styles['Normal']
        norm_style.paragraph_format.line_spacing = 1

        custom_style = doc.styles.add_style('CustomStyle', WD_STYLE_TYPE.CHARACTER)
        custom_style.font.size = Pt(12)
        custom_style.font.name = 'Georgia'
        
        custom_style_2 = doc.styles.add_style('CustomStyle2', WD_STYLE_TYPE.CHARACTER)
        custom_style_2.font.size = Pt(11.5)
        custom_style_2.font.name = 'Georgia'

        list_style = doc.styles['Bullet New']
        list_style.paragraph_format.line_spacing = 1

        # Add scores 
        for i, paragraph in enumerate(doc.paragraphs):
            if "Scores are reported here as standard scores" in paragraph.text:
                total = 0
                for test in check_scores:
                    if check_scores[test]:
                        total += 1 
                        if total == 1:
                            paragraph.insert_paragraph_before()
                            paragraph.insert_paragraph_before().add_run("Psychoeducational Testing:", style='CustomStyle').font.underline = True
                        add_score(paragraph, score_data=scores[test])

                if total == 0:
                    delete_paragraph(paragraph)
            
            if "[[Recommendations]]" in paragraph.text:
                for rec, checked in check_rec.items():
                    if checked:
                        func = globals().get(f"add_{rec}")
                        if callable(func):
                            func(paragraph)
                
                delete_paragraph(paragraph)
                
            if "[[Behavioral Presentation]]" in paragraph.text:
                add_behavior_presentation(paragraph, st.session_state.behavior_observation_mod12_no_autism)
            
            if "[[District Grade School Setting]]" in paragraph.text:
                add_school(paragraph)
        
        # Edit document
        for word in replace_word:
            docxedit.replace_string(doc, old_string=word, new_string=replace_word[word])

        # Replace for lists separated by comma:
        for word in comma:
            new_word = ", ".join(comma[word])
            docxedit.replace_string(doc, old_string=word, new_string=new_word)

        # Replace for lists separated by new line:
        for word in lines:
            new_word = "\n".join(lines[word])
            docxedit.replace_string(doc, old_string=word, new_string=new_word)

        # Save content to file
        doc.save(filename)

        # Replace for lists separated by bullet points
        tpl=DocxTemplate(filename)
        print("Load template!")

        tpl.render(bullet)
        print("Bullet rendered!")

        tpl.save(filename)
        print("File saved at", filename)

        # Download 
        bio = io.BytesIO()
        document = Document(filename)
        document.save(bio)
        
        st.download_button(
            label="Click here to download",
            key="report_download",
            data=bio.getvalue(),
            file_name=filename,
            mime="docx"
        )