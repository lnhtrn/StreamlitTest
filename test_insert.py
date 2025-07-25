import streamlit as st
from docx import Document
import yaml
import io
import docxedit
import datetime
import re 
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE, WD_STYLE
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH

VINELAND = True

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def get_ordinal(number):
    suffix = 'th' if 11 <= int(number) <= 13 else {"1": 'st', "2": 'nd', "3": 'rd'}.get(number[-1], 'th')
    return suffix

def replace_with_superscript(para, old_text, number_part):
    superscript_part = get_ordinal(number_part)
    if old_text in para.text:
        # Save everything before, match, and after
        before, match, after = para.text.partition(old_text)
        
        # Clear existing runs
        para.clear()
        
        # Add before text
        para.add_run(before, style='CustomStyle')
        
        # Add number part
        para.add_run(number_part, style='CustomStyle')
        
        # Add superscript part
        sup_run = para.add_run(superscript_part, style='CustomStyle')
        sup_run.font.superscript = True
        
        # Add after text
        para.add_run(after, style='CustomStyle')


def replace_ordinal_with_superscript(para, full_text):
    # Regex to find ordinal numbers like 1st, 2nd, 77th, 103rd, etc.
    pattern = re.compile(r'(\d+)(st|nd|rd|th)')
    paragraph = para.insert_paragraph_before()  # Remove all existing runs

    last_index = 0
    for match in pattern.finditer(full_text):
        # Add text before the match
        paragraph.add_run(full_text[last_index:match.start()], style='CustomStyle')

        # Add number part (e.g., "77")
        paragraph.add_run(match.group(1), style='CustomStyle')

        # Add superscript suffix (e.g., "th")
        sup_run = paragraph.add_run(match.group(2), style='CustomStyle')
        sup_run.font.superscript = True

        last_index = match.end()

    # Add remaining text after last match
    paragraph.add_run(full_text[last_index:], style='CustomStyle')
    delete_paragraph(para)

vineland_score = {
    # columns = 2
    "[[Adaptive Behavior Composite]]": "35",
    "[[Adaptive Behavior Composite Percentile]]": 1,
    # columns = 2
    "[[Communication]]": "34",
    "[[Communication Percentile]]": 1,
    # columns = 3
    "[[Receptive]]": "2:8",
    "[[Expressive]]": "3:7",
    "[[Written]]": "22:0+",
    # columns = 2
    "[[Daily Living Skills]]": "49*",
    "[[Daily Living Skills Percentile]]": 1,
    # columns = 3
    "[[Personal]]": "7:10",
    "[[Domestic/Numeric]]": "13:9",
    "[[Community]]": "17:6",
    # columns = 2
    "[[Socialization]]": "20",
    "[[Socialization Percentile]]": 1,
    # columns = 3
    "[[Interpersonal Relationships]]": "1:2",
    "[[Play & Leisure]]": "2:3",
    "[[Coping Skills]]": "4:6",
}

doc = Document('templates/template_mod_4.docx')
if doc:
    ### create document style
    doc_style = doc.styles['Normal']
    font = doc_style.font
    font.name = 'Georgia'
    font.size = Pt(12)

    custom_style = doc.styles.add_style('CustomStyle', WD_STYLE_TYPE.CHARACTER)
    custom_style.font.size = Pt(12)
    custom_style.font.name = 'Georgia'

    # If we have Vineland Score
    if VINELAND:
        for word in vineland_score:
            docxedit.replace_string(doc, old_string=word, new_string=vineland_score[word])

        for paragraph in doc.paragraphs:
            if "[[Vineland Analysis]]" in paragraph.text:
                # Add page break
                paragraph.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)

                p = paragraph.insert_paragraph_before()
                r = p.add_run("Interpretation of VABS-3 Results – Informant Report", style='CustomStyle')
                r.bold = True
                r.italic = True

                paragraph.insert_paragraph_before().add_run("\nInterpretation of VABS-3 Results – Informant Report", style='CustomStyle')


    # If no Vineland Score
    else:
        allTables = doc.tables

        for activeTable in allTables:
            if activeTable.cell(0,0).paragraphs[0].text == 'Adaptive Behavior Composite':
                activeTable._element.getparent().remove(activeTable._element)

        # Test paragraph 
        for i, paragraph in enumerate(doc.paragraphs):
            # Vineland Informant Report 
            if "[[Vineland_Start]]" in paragraph.text:
                vineland_start = i 

        if vineland_start:
            print(vineland_start)
            for index in range(vineland_start, vineland_start+4, 1):
                try:
                    delete_paragraph(doc.paragraphs[index])
                except:
                    print("Out of range at index", i)
        else:
            print("Cannot find Vineland Start")

        
        for i, paragraph in enumerate(doc.paragraphs):
            if "The VABS-3 yields information about an individual’s adaptive functioning" in paragraph.text:
                delete_paragraph(paragraph)
            if "[[Vineland Analysis]]" in paragraph.text:
                delete_paragraph(paragraph)

        # if "[[WAIS-Analysis]]" in paragraph.text:
        #     replace_ordinal_with_superscript(paragraph, wais_analysis)

        # if "[[Test Percentile]]" in paragraph.text:
        #     replace_with_superscript(paragraph, "[[Test Percentile]]", "50")

doc.save("other_misc/test_delete_table.docx")